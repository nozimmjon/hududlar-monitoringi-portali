from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "data_source" / "Кафолат хатлар имзога" / "2. Андижон"
DATA_JSON = ROOT / "platform prototypes" / "andijon_full_pilot_assets" / "andijon_full_pilot_data.json"
OUT_HTML = ROOT / "platform prototypes" / "andijon_commitment_workflow_v4.html"


def read_paragraphs() -> list[str]:
    doc = Document(SRC / "0. Кафолат хати (Андижон).docx")
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def p(paragraphs: list[str], *numbers: int) -> list[dict[str, str | int]]:
    return [{"n": n, "text": paragraphs[n - 1]} for n in numbers if 0 < n <= len(paragraphs)]


def build_commitments(paragraphs: list[str]) -> list[dict]:
    return [
        {
            "id": "M-01",
            "section": "I. Макро",
            "title": "ЯҲМ ўсишини таъминлаш",
            "kind": "KPI",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 3),
            "excel": ["1.1-1.5-жадваллар (макро).xlsx · 1.1. ЯҲМ"],
            "kpis": ["ЯҲМ"],
            "districtKey": None,
            "taskParagraphs": [],
        },
        {
            "id": "M-02",
            "section": "I. Макро",
            "title": "Саноат маҳсулотлари ва ҳудудий саноат",
            "kind": "KPI + топшириқ",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 5, 7),
            "excel": [
                "1.1-1.5-жадваллар (макро).xlsx · 1.1. ЯҲМ",
                "1.1-1.5-жадваллар (макро).xlsx · 1.2. Саноат",
                "1.1-1.5-жадваллар (макро).xlsx · 1.3. Ҳудудий саноат",
            ],
            "kpis": ["Саноат маҳсулотлари", "Махсус иқтисодий ва саноат зоналари", "Энергия тежаш"],
            "districtKey": "industry",
            "taskParagraphs": p(paragraphs, 9, 12, 13, 15, 16, 19),
        },
        {
            "id": "M-03",
            "section": "I. Макро",
            "title": "Бозор хизматлари ҳажмини ошириш",
            "kind": "KPI + топшириқ",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 20),
            "excel": ["1.1-1.5-жадваллар (макро).xlsx · 1.5. Бозор хизматлари"],
            "kpis": ["Бозор хизматлари"],
            "districtKey": "services",
            "taskParagraphs": p(paragraphs, 22, 23, 24),
        },
        {
            "id": "M-04",
            "section": "I. Макро",
            "title": "Қишлоқ хўжалиги маҳсулотлари",
            "kind": "KPI + топшириқ",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 25),
            "excel": ["1.1-1.5-жадваллар (макро).xlsx · 1.4. ҚХ"],
            "kpis": ["Қишлоқ хўжалиги маҳсулотлари"],
            "districtKey": "agriculture",
            "taskParagraphs": p(paragraphs, 27),
        },
        {
            "id": "M-05",
            "section": "I. Макро",
            "title": "Қурилиш ишлари ва уй-жойлар",
            "kind": "KPI + топшириқ",
            "link": "Қисман",
            "paragraphs": p(paragraphs, 28, 29),
            "excel": ["1.1-1.5-жадваллар (макро).xlsx · 1.1. ЯҲМ"],
            "kpis": ["Қурилиш ишлари"],
            "districtKey": None,
            "taskParagraphs": p(paragraphs, 29),
        },
        {
            "id": "M-06",
            "section": "II. Инфляция",
            "title": "Инфляция ва озиқ-овқат барқарорлиги",
            "kind": "Топшириқ + баланс",
            "link": "Қисман",
            "paragraphs": p(paragraphs, 31, 38, 39, 40, 43, 45, 46, 47),
            "excel": [
                "2.1-2.2-жадваллар (инфляция).xlsx · 1.1. Баланс",
                "2.1-2.2-жадваллар (инфляция).xlsx · 1.2. Омборлар",
            ],
            "kpis": ["Инфляция чегараси", "Озиқ-овқат баланси", "Омбор сиғими"],
            "districtKey": "warehouses",
            "taskParagraphs": p(paragraphs, 32, 33, 34, 35, 36, 37, 41, 42),
        },
        {
            "id": "M-07",
            "section": "III. Бюджет",
            "title": "Бюджет тушумлари прогнозини бажариш",
            "kind": "KPI + чора",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 49, 56),
            "excel": ["3-жадвал (бюджет).xlsx · тушум"],
            "kpis": ["Бюджет тушумлари", "Қўшимча тушум"],
            "districtKey": "budget",
            "taskParagraphs": p(paragraphs, 51, 52, 53, 54, 55),
        },
        {
            "id": "M-08",
            "section": "IV. Инвестиция",
            "title": "Бюджет маблағлари ҳисобидан лойиҳалар",
            "kind": "KPI + объект",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 59, 63, 65),
            "excel": ["4.1-жадвал (бюджет инвестка).xlsx · 2.Анд"],
            "kpis": ["Ўзлаштириш", "Фойдаланишга топшириш"],
            "districtKey": "budget_investment",
            "taskParagraphs": p(paragraphs, 60, 61, 62, 64, 66),
        },
        {
            "id": "M-09",
            "section": "IV. Инвестиция",
            "title": "Хорижий инвестициялар ва лойиҳалар",
            "kind": "KPI + лойиҳа",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 68, 70, 71, 72, 73),
            "excel": ["4.2-жадвал (инвестициялар).xlsx · 4,2-хорижий инв"],
            "kpis": ["Хорижий инвестициялар", "Лойиҳа сони", "Иш ўрни"],
            "districtKey": "foreign_investment",
            "taskParagraphs": p(paragraphs, 69, 74),
        },
        {
            "id": "M-10",
            "section": "V. Экспорт",
            "title": "Экспорт ҳажми, экспортчилар ва бозорлар",
            "kind": "KPI + манзилли иш",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 76, 78, 81, 82, 87, 88),
            "excel": [
                "5.1-5.2-жадваллар (экспорт).xlsx · 5-жадвал",
                "5.1-5.2-жадваллар (экспорт).xlsx · 02_Анд",
                "5.1-5.2-жадваллар (экспорт).xlsx · Корхона сони",
            ],
            "kpis": ["Экспорт ҳажми", "Экспортчи корхоналар", "Экспорт ўсиши"],
            "districtKey": "export",
            "taskParagraphs": p(paragraphs, 79, 80, 85, 86, 87, 89),
        },
        {
            "id": "M-11",
            "section": "VI. Бандлик",
            "title": "Ишсизликни камайтириш ва иш ўринлари",
            "kind": "KPI + топшириқ",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 91, 92, 93, 94, 99),
            "excel": ["6-жадвал (бандлик ва камбағаллик даражаси).xlsx · 6. Камбағаллик"],
            "kpis": ["Ишсизлик", "Доимий ишга жойлаштириш", "Легаллаштириш", "Микролойиҳалар"],
            "districtKey": "employment",
            "taskParagraphs": p(paragraphs, 95, 96, 97, 98, 99),
        },
        {
            "id": "M-12",
            "section": "VI. Камбағаллик",
            "title": "Камбағалликни қисқартириш ва холи маҳаллалар",
            "kind": "KPI + хизмат",
            "link": "Тўлиқ",
            "paragraphs": p(paragraphs, 101, 102, 103, 104, 105),
            "excel": ["6-жадвал (бандлик ва камбағаллик даражаси).xlsx · 6. Камбағаллик"],
            "kpis": ["Камбағаллик", "Холи МФЙ", "Индивидуал режа хизматлари"],
            "districtKey": "employment",
            "taskParagraphs": p(paragraphs, 106, 107, 108),
        },
    ]


def main() -> None:
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    paragraphs = read_paragraphs()
    payload = {
        "meta": data["meta"],
        "regional": data["regional"],
        "districts": data["districts"],
        "sources": data["sources"],
        "data_quality": data["data_quality"],
        "commitments": build_commitments(paragraphs),
    }
    html = HTML.replace("__DATA__", json.dumps(payload, ensure_ascii=False).replace("</", "<\\/"))
    OUT_HTML.write_text(html, encoding="utf-8")
    print(OUT_HTML)


HTML = r"""<!doctype html>
<html lang="uz-Cyrl">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Кафолат хати workflow · Андижон</title>
  <style>
    :root {
      --bg: #f4f2ec;
      --panel: #fffdf8;
      --ink: #192230;
      --muted: #626b78;
      --line: #d7d1c5;
      --navy: #173861;
      --blue: #2158a8;
      --green: #16724a;
      --amber: #9d6508;
      --red: #b43a38;
      --grey: #7a818b;
      --shadow: 0 18px 44px rgba(36, 41, 53, .08);
      --r: 8px;
      font-family: "Aptos", "Segoe UI", sans-serif;
    }

    * { box-sizing: border-box; }
    body {
      margin: 0;
      color: var(--ink);
      background:
        linear-gradient(90deg, rgba(23, 56, 97, .055) 1px, transparent 1px) 0 0 / 72px 72px,
        var(--bg);
    }

    button, input, select { font: inherit; }
    button { cursor: pointer; }

    .app {
      max-width: 1520px;
      margin: 0 auto;
      padding: 18px 22px 36px;
    }

    .hero {
      background: var(--navy);
      color: #f8fbff;
      border-radius: var(--r);
      box-shadow: var(--shadow);
      overflow: hidden;
    }

    .hero-main {
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      gap: 24px;
      align-items: end;
      padding: 24px;
      border-bottom: 1px solid rgba(255,255,255,.16);
    }

    .eyebrow {
      color: #bad0e9;
      font-size: 12px;
      text-transform: uppercase;
      letter-spacing: .04em;
      font-weight: 800;
    }

    h1, h2, h3 { margin: 0; letter-spacing: 0; }
    h1 { margin-top: 6px; font-size: 30px; line-height: 1.08; }
    .hero p { margin: 8px 0 0; color: #dce8f5; }

    .controls {
      display: flex;
      gap: 10px;
      flex-wrap: wrap;
      justify-content: flex-end;
    }

    select, input {
      border: 1px solid rgba(255,255,255,.34);
      background: rgba(255,255,255,.08);
      color: inherit;
      border-radius: 6px;
      min-height: 40px;
      padding: 9px 12px;
      outline: none;
    }

    input::placeholder { color: rgba(255,255,255,.68); }

    .workflow {
      display: grid;
      grid-template-columns: repeat(6, minmax(0, 1fr));
      border-bottom: 1px solid rgba(255,255,255,.16);
    }

    .step {
      padding: 15px 16px;
      border-right: 1px solid rgba(255,255,255,.16);
      min-height: 86px;
    }

    .step:last-child { border-right: 0; }
    .step span { display: block; color: #bad0e9; font-size: 12px; }
    .step strong { display: block; margin-top: 5px; font-size: 15px; line-height: 1.25; }

    .nav {
      display: flex;
      overflow-x: auto;
      background: rgba(255,255,255,.07);
    }

    .nav button {
      border: 0;
      border-right: 1px solid rgba(255,255,255,.14);
      background: transparent;
      color: #dbe8f7;
      padding: 13px 18px;
      white-space: nowrap;
      font-weight: 800;
    }

    .nav button.active { background: var(--panel); color: var(--navy); }

    .view { display: none; padding-top: 22px; }
    .view.active { display: block; }

    .layout {
      display: grid;
      grid-template-columns: 420px minmax(0, 1fr);
      gap: 22px;
      align-items: start;
    }

    .panel {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: var(--r);
      box-shadow: var(--shadow);
      overflow: hidden;
    }

    .panel-head {
      padding: 16px 18px;
      border-bottom: 1px solid var(--line);
      background: #f3f0e8;
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 14px;
    }

    .panel-head h2 { font-size: 20px; }
    .panel-head p { margin: 5px 0 0; color: var(--muted); font-size: 13px; line-height: 1.35; }

    .commit-list {
      max-height: calc(100vh - 278px);
      overflow: auto;
    }

    .commit-btn {
      display: grid;
      grid-template-columns: 56px minmax(0, 1fr);
      gap: 12px;
      width: 100%;
      border: 0;
      border-bottom: 1px solid var(--line);
      background: transparent;
      color: var(--ink);
      text-align: left;
      padding: 13px 16px;
    }

    .commit-btn:hover { background: #faf7ef; }
    .commit-btn.active { background: #eef3fb; box-shadow: inset 4px 0 0 var(--blue); }
    .code {
      font-weight: 900;
      color: var(--blue);
      font-size: 13px;
      padding-top: 2px;
    }
    .commit-btn strong { display: block; font-size: 14px; line-height: 1.25; }
    .meta { display: block; color: var(--muted); font-size: 12px; margin-top: 5px; line-height: 1.35; }

    .badges { display: flex; gap: 6px; flex-wrap: wrap; margin-top: 8px; }
    .badge {
      display: inline-flex;
      align-items: center;
      border-radius: 999px;
      padding: 3px 8px;
      font-size: 11px;
      font-weight: 900;
      border: 1px solid transparent;
      white-space: nowrap;
    }
    .full { color: var(--green); background: #e5f3eb; border-color: #badfc9; }
    .partial { color: var(--amber); background: #fff1d2; border-color: #e6c57a; }
    .task { color: var(--blue); background: #e7eef9; border-color: #c4d6f1; }
    .missing { color: var(--grey); background: #efede7; border-color: #dad5ca; }
    .red { color: var(--red); background: #fde7e5; border-color: #efb9b6; }

    .detail-grid {
      display: grid;
      grid-template-columns: minmax(0, 1fr) 340px;
      gap: 18px;
      padding: 18px;
    }

    .story {
      display: grid;
      gap: 14px;
    }

    .mini-title {
      font-size: 12px;
      color: var(--muted);
      font-weight: 900;
      text-transform: uppercase;
      letter-spacing: .03em;
      margin-bottom: 8px;
    }

    .quote {
      border-left: 4px solid var(--blue);
      background: #f8f6ef;
      padding: 12px 14px;
      font-size: 13px;
      line-height: 1.5;
    }

    .quote b { color: var(--blue); }

    .source-list, .task-list, .kpi-list {
      display: grid;
      gap: 8px;
    }

    .source-item, .task-item, .kpi-item {
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 10px 11px;
      background: #fffaf1;
      font-size: 13px;
      line-height: 1.35;
    }

    .sidebox {
      border: 1px solid var(--line);
      border-radius: 6px;
      overflow: hidden;
      background: #fffaf1;
    }

    .sidebox .row {
      display: grid;
      grid-template-columns: 1fr auto;
      gap: 10px;
      padding: 11px 12px;
      border-bottom: 1px solid var(--line);
      font-size: 13px;
    }

    .sidebox .row:last-child { border-bottom: 0; }
    .sidebox span { color: var(--muted); }
    .sidebox strong { text-align: right; }

    table {
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
    }

    th {
      text-align: left;
      background: #f2efe7;
      color: #495363;
      padding: 10px 12px;
      border-bottom: 1px solid var(--line);
      white-space: nowrap;
    }

    td {
      padding: 11px 12px;
      border-bottom: 1px solid #e8e2d7;
      vertical-align: top;
    }

    tr:last-child td { border-bottom: 0; }
    .num { font-weight: 900; white-space: nowrap; }
    .sub { display: block; color: var(--muted); font-size: 12px; margin-top: 3px; line-height: 1.35; }
    .table-wrap { overflow-x: auto; }

    .two-col {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 22px;
    }

    .stats {
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      border-bottom: 1px solid var(--line);
    }

    .stat {
      padding: 14px 16px;
      border-right: 1px solid var(--line);
      background: #fbf8f0;
    }

    .stat:last-child { border-right: 0; }
    .stat span { color: var(--muted); font-size: 12px; display: block; }
    .stat strong { display: block; margin-top: 5px; font-size: 24px; }

    .empty {
      padding: 28px;
      text-align: center;
      color: var(--muted);
    }

    @media (max-width: 1120px) {
      .hero-main, .layout, .detail-grid, .two-col { grid-template-columns: 1fr; }
      .workflow, .stats { grid-template-columns: repeat(2, minmax(0, 1fr)); }
      .controls { justify-content: flex-start; }
      .commit-list { max-height: none; }
    }

    @media (max-width: 680px) {
      .app { padding: 10px; }
      .hero-main { padding: 16px; }
      h1 { font-size: 23px; }
      .controls { display: grid; grid-template-columns: 1fr; }
      select, input { width: 100%; }
      .workflow { grid-template-columns: repeat(2, minmax(0, 1fr)); }
      .stats { grid-template-columns: 1fr; }
      .step {
        min-height: 58px;
        padding: 10px 12px;
        border-bottom: 1px solid rgba(255,255,255,.16);
      }
      .step strong { font-size: 13px; }
      .stat { border-right: 0; border-bottom: 1px solid rgba(255,255,255,.16); }
      .stat { border-bottom-color: var(--line); }
      table { min-width: 780px; }
    }
  </style>
</head>
<body>
  <div class="app">
    <header class="hero">
      <div class="hero-main">
        <div>
          <div class="eyebrow">Андижон пилоти · кафолат хати билан жадваллар интеграцияси</div>
          <h1>Мажбуриятдан ижрогача мониторинг workflow</h1>
          <p>Ҳар бир ваъда матний асос, рақамли асос, туман кесими, топшириқ ва далил билан боғланади.</p>
        </div>
        <div class="controls">
          <select id="sectionFilter"></select>
          <input id="search" type="search" placeholder="Мажбурият, KPI, туман ёки манба">
        </div>
      </div>
      <div class="workflow">
        <div class="step"><span>1-қадам</span><strong>Кафолат хати мажбурияти</strong></div>
        <div class="step"><span>2-қадам</span><strong>Excel рақамли асос</strong></div>
        <div class="step"><span>3-қадам</span><strong>Давр ва туман кесими</strong></div>
        <div class="step"><span>4-қадам</span><strong>Ижро ҳолати</strong></div>
        <div class="step"><span>5-қадам</span><strong>Топшириқ ва масъул</strong></div>
        <div class="step"><span>6-қадам</span><strong>Далил ва қарор</strong></div>
      </div>
      <nav class="nav" id="nav"></nav>
    </header>

    <section id="view-map" class="view active"></section>
    <section id="view-matrix" class="view"></section>
    <section id="view-districts" class="view"></section>
    <section id="view-tasks" class="view"></section>
    <section id="view-sources" class="view"></section>
  </div>

  <script>
    const DATA = __DATA__;
    const NAV = [
      ["map", "Мажбуриятлар харитаси"],
      ["matrix", "Боғланиш матрицаси"],
      ["districts", "Туман кесими"],
      ["tasks", "Топшириқлар"],
      ["sources", "Манбалар аудити"]
    ];
    const state = { view: "map", selected: DATA.commitments[0].id, section: "all", search: "" };
    const $ = (s) => document.querySelector(s);
    const esc = (v) => String(v ?? "").replace(/[&<>"']/g, (ch) => ({ "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;", "'": "&#39;" }[ch]));

    function fmt(value, digits = 1) {
      if (value === null || value === undefined || value === "" || Number.isNaN(Number(value))) return "—";
      return Number(value).toLocaleString("uz-UZ", { maximumFractionDigits: digits, minimumFractionDigits: 0 });
    }

    function badgeClass(c) {
      if (c.link === "Тўлиқ") return "full";
      if (c.link === "Қисман") return "partial";
      return "missing";
    }

    function filteredCommitments() {
      const q = state.search.trim().toLowerCase();
      return DATA.commitments.filter((c) => {
        const sectionOk = state.section === "all" || c.section === state.section;
        const hay = `${c.id} ${c.section} ${c.title} ${c.kind} ${c.link} ${c.kpis.join(" ")} ${c.excel.join(" ")}`.toLowerCase();
        return sectionOk && (!q || hay.includes(q));
      });
    }

    function selectedCommitment() {
      return DATA.commitments.find((c) => c.id === state.selected) || filteredCommitments()[0] || DATA.commitments[0];
    }

    function renderNav() {
      $("#nav").innerHTML = NAV.map(([id, label]) => `<button type="button" data-view="${id}" class="${state.view === id ? "active" : ""}">${label}</button>`).join("");
      document.querySelectorAll("[data-view]").forEach((btn) => btn.addEventListener("click", () => {
        state.view = btn.dataset.view;
        render();
      }));
    }

    function renderFilters() {
      const sections = ["all", ...new Set(DATA.commitments.map((c) => c.section))];
      $("#sectionFilter").innerHTML = sections.map((s) => `<option value="${esc(s)}">${s === "all" ? "Барча бўлимлар" : esc(s)}</option>`).join("");
      $("#sectionFilter").value = state.section;
    }

    function commitmentList() {
      const rows = filteredCommitments();
      return `<div class="commit-list">${rows.map((c) => `
        <button type="button" class="commit-btn ${c.id === selectedCommitment().id ? "active" : ""}" data-commit="${c.id}">
          <span class="code">${esc(c.id)}</span>
          <span>
            <strong>${esc(c.title)}</strong>
            <span class="meta">${esc(c.section)} · ${esc(c.kind)}</span>
            <span class="badges">
              <span class="badge ${badgeClass(c)}">${esc(c.link)} боғланган</span>
              <span class="badge task">${c.taskParagraphs.length} топшириқ</span>
            </span>
          </span>
        </button>`).join("") || `<div class="empty">Мажбурият топилмади.</div>`}</div>`;
    }

    function bindCommitButtons() {
      document.querySelectorAll("[data-commit]").forEach((btn) => btn.addEventListener("click", () => {
        state.selected = btn.dataset.commit;
        render();
      }));
    }

    function selectedKpiRows(c) {
      const r = DATA.regional;
      if (c.id === "M-01") {
        const x = r.macro.find((i) => i.indicator === "ЯҲМ");
        return [["ЯҲМ", "I чорак амалда", `${fmt(x.q1_value)} млрд сўм / ${fmt(x.q1_growth - 100)}%`, x.source],
          ["ЯҲМ", "I ярим йиллик прогноз", `${fmt(x.h1_value)} млрд сўм / ${fmt(x.h1_growth - 100)}%`, x.source],
          ["ЯҲМ", "Йил якуни прогноз", `${fmt(x.year_value)} млрд сўм / ${fmt(x.year_growth - 100)}%`, x.source]];
      }
      if (c.districtKey && ["industry", "agriculture", "services"].includes(c.districtKey)) {
        const label = { industry: "Саноат", agriculture: "Қишлоқ хўжалиги", services: "Бозор хизматлари" }[c.districtKey];
        const x = DATA.districts[0].data[c.districtKey];
        const total = r.macro.find((i) => i.indicator.includes(label) || i.indicator === "Саноат маҳсулотлари");
        return total ? [[label, "I чорак амалда", `${fmt(total.q1_value)} млрд сўм / ${fmt(total.q1_growth - 100)}%`, total.source],
          [label, "I ярим йиллик прогноз", `${fmt(total.h1_value)} млрд сўм / ${fmt(total.h1_growth - 100)}%`, total.source],
          [label, "Йил якуни прогноз", `${fmt(total.year_value)} млрд сўм / ${fmt(total.year_growth - 100)}%`, total.source]] : [];
      }
      if (c.id === "M-06") {
        return [["Инфляция", "II чорак", "2,9%дан оширмаслик", "0. Кафолат хати (Андижон).docx"],
          ["Инфляция", "Йил якуни", "6,6%дан оширмаслик", "0. Кафолат хати (Андижон).docx"],
          ["Омборлар", "Йил якуни", "409 та омбор / 145,6 минг тонна", "2.1-2.2-жадваллар · 1.2. Омборлар"]];
      }
      if (c.id === "M-07") return [["Бюджет", "I ярим йиллик", `${fmt(r.budget.h1_expected)} / ${fmt(r.budget.h1_plan)} млрд сўм · ${fmt(r.budget.h1_execution_pct)}%`, r.budget.source], ["Бюджет", "Йил якуни", `${fmt(r.budget.year_expected)} / ${fmt(r.budget.year_plan)} млрд сўм`, r.budget.source]];
      if (c.id === "M-08") return [["Ўзлаштириш", "I ярим йиллик", `${fmt(r.budget_investment.h1_absorption)} млн сўм · ${fmt(r.budget_investment.h1_pct)}%`, r.budget_investment.source], ["Объект", "Йил якуни", `${fmt(r.budget_investment.commissioning_year_count, 0)} та объект`, r.budget_investment.source]];
      if (c.id === "M-09") return [["Хорижий инвестиция", "I чорак", `${fmt(r.foreign_investment.q1_actual)} млн $`, r.foreign_investment.source], ["Хорижий инвестиция", "I ярим йиллик", `${fmt(r.foreign_investment.h1_expected)} млн $`, r.foreign_investment.source], ["Лойиҳалар", "Йил якуни", `${fmt(r.foreign_investment.year_expected)} млн $ / ${fmt(r.foreign_investment.h1_projects, 0)} лойиҳа H1`, r.foreign_investment.source]];
      if (c.id === "M-10") return [["Экспорт", "I чорак", `${fmt(r.export.q1_value / 1000)} млн $ / ${fmt(r.export.q1_growth - 100)}%`, r.export.source], ["Экспорт", "I ярим йиллик", `${fmt(r.export.h1_expected / 1000)} млн $ / ${fmt(r.export.h1_growth - 100)}%`, r.export.source], ["Экспорт", "Йил якуни", `${fmt(r.export.year_expected / 1000)} млн $ / ${fmt(r.export.year_growth - 100)}%`, r.export.source]];
      if (c.id === "M-11") return [["Ишсизлик", "I ярим йиллик", `${fmt(r.employment.unemployment_h1)}%`, r.employment.source], ["Ишсизлик", "Йил якуни", `${fmt(r.employment.unemployment_year)}%`, r.employment.source], ["Доимий иш", "Йил якуни", `${fmt(r.employment.jobs_year, 1)} минг нафар`, r.employment.source]];
      if (c.id === "M-12") return [["Камбағаллик", "I ярим йиллик", `${fmt(r.employment.poverty_h1)}%`, r.employment.source], ["Камбағаллик", "Йил якуни", `${fmt(r.employment.poverty_year)}%`, r.employment.source], ["Холи МФЙ", "Йил якуни", `${fmt(r.employment.mfy_year, 0)} та`, r.employment.source]];
      return [];
    }

    function detail(c) {
      const kpiRows = selectedKpiRows(c);
      return `<div class="panel">
        <div class="panel-head">
          <div>
            <h2>${esc(c.id)} · ${esc(c.title)}</h2>
            <p>${esc(c.section)} · ${esc(c.kind)} · кафолат хати ва жадвал боғланиши</p>
          </div>
          <span class="badge ${badgeClass(c)}">${esc(c.link)}</span>
        </div>
        <div class="detail-grid">
          <div class="story">
            <div>
              <div class="mini-title">Матний асос: кафолат хати</div>
              ${c.paragraphs.map((x) => `<div class="quote"><b>${x.n}-параграф.</b> ${esc(x.text)}</div>`).join("")}
            </div>
            <div>
              <div class="mini-title">Рақамли асос: Excel жадваллар</div>
              <div class="source-list">${c.excel.map((x) => `<div class="source-item">${esc(x)}</div>`).join("")}</div>
            </div>
            <div>
              <div class="mini-title">Кўрсаткичлар</div>
              <div class="table-wrap"><table>
                <thead><tr><th>KPI</th><th>Давр</th><th>Рақам</th><th>Манба</th></tr></thead>
                <tbody>${kpiRows.map((r) => `<tr><td class="num">${esc(r[0])}</td><td>${esc(r[1])}</td><td class="num">${esc(r[2])}</td><td><span class="sub">${esc(r[3])}</span></td></tr>`).join("") || `<tr><td colspan="4">Бу мажбурият асосан топшириқ/далил орқали мониторинг қилинади.</td></tr>`}</tbody>
              </table></div>
            </div>
          </div>
          <aside>
            <div class="sidebox">
              <div class="row"><span>Матний параграф</span><strong>${c.paragraphs.length}</strong></div>
              <div class="row"><span>Excel манба</span><strong>${c.excel.length}</strong></div>
              <div class="row"><span>KPI</span><strong>${c.kpis.length}</strong></div>
              <div class="row"><span>Топшириқ</span><strong>${c.taskParagraphs.length}</strong></div>
              <div class="row"><span>Туман кесими</span><strong>${c.districtKey ? "бор" : "йўқ/қисман"}</strong></div>
            </div>
            <div style="height: 12px;"></div>
            <div class="mini-title">Топшириққа айланадиган матнлар</div>
            <div class="task-list">${c.taskParagraphs.map((x) => `<div class="task-item"><b>${x.n}-п.</b> ${esc(x.text)}</div>`).join("") || `<div class="task-item">Алоҳида топшириқ ажратилмаган.</div>`}</div>
          </aside>
        </div>
      </div>`;
    }

    function renderMap() {
      const c = selectedCommitment();
      $("#view-map").innerHTML = `<div class="layout">
        <aside class="panel">
          <div class="panel-head"><div><h2>Кафолат мажбуриятлари</h2><p>Ҳар бир қатор матн ва рақамли иловага боғланган.</p></div></div>
          ${commitmentList()}
        </aside>
        ${detail(c)}
      </div>`;
      bindCommitButtons();
    }

    function renderMatrix() {
      $("#view-matrix").innerHTML = `<section class="panel">
        <div class="panel-head"><div><h2>Кафолат хати ↔ жадваллар боғланиш матрицаси</h2><p>Платформа ядроси: мажбурият марказда, рақамлар ва топшириқлар унга уланади.</p></div></div>
        <div class="table-wrap"><table>
          <thead><tr><th>ID</th><th>Бўлим</th><th>Мажбурият</th><th>Кафолат хати</th><th>Excel асос</th><th>KPI/топшириқ</th><th>Туман кесими</th><th>Боғланиш</th></tr></thead>
          <tbody>${filteredCommitments().map((c) => `<tr>
            <td class="num">${esc(c.id)}</td><td>${esc(c.section)}</td><td class="num">${esc(c.title)}</td>
            <td>${c.paragraphs.map((x) => `${x.n}-п.`).join(", ")}</td>
            <td>${c.excel.map((x) => `<span class="sub">${esc(x)}</span>`).join("")}</td>
            <td>${c.kpis.map((x) => `<span class="sub">${esc(x)}</span>`).join("")}</td>
            <td>${c.districtKey ? "бор" : "йўқ/қисман"}</td>
            <td><span class="badge ${badgeClass(c)}">${esc(c.link)}</span></td>
          </tr>`).join("")}</tbody>
        </table></div>
      </section>`;
    }

    function districtValue(d, key) {
      const data = d.data[key];
      if (!data) return ["—", "—", "—"];
      if (["industry", "agriculture", "services"].includes(key)) return [`${fmt(data.q1_value)} / +${fmt(data.q1_growth - 100)}%`, `${fmt(data.h1_value)} / +${fmt(data.h1_growth - 100)}%`, `${fmt(data.year_value)} / +${fmt(data.year_growth - 100)}%`];
      if (key === "budget") return ["—", `${fmt(data.h1_expected)} / ${fmt(data.h1_execution_pct)}%`, `${fmt(data.year_expected)}`];
      if (key === "budget_investment") return [`${fmt(data.q1_absorption)} / ${fmt(data.q1_pct)}%`, `${fmt(data.h1_absorption)} / ${fmt(data.h1_pct)}%`, `${fmt(data.year_absorption)} / ${fmt(data.year_pct)}%`];
      if (key === "foreign_investment") return [`${fmt(data.q1_actual)} / ${fmt(data.q1_pct)}`, `${fmt(data.h1_expected)} / ${fmt(data.h1_pct)}`, `${fmt(data.year_expected)} / ${fmt(data.year_pct)}`];
      if (key === "export") return [`${fmt(data.q1_value / 1000)} / +${fmt(data.q1_growth - 100)}%`, `${fmt(data.h1_expected / 1000)} / +${fmt(data.h1_growth - 100)}%`, `${fmt(data.year_expected / 1000)} / +${fmt(data.year_growth - 100)}%`];
      if (key === "employment") return [`ишсизлик ${fmt(data.unemployment_h1)}%`, `камбағаллик ${fmt(data.poverty_h1)}%`, `иш ўрни ${fmt(data.jobs_year)} минг`];
      if (key === "warehouses") return [`захира ${fmt(data.reserve_warehouses,0)} та`, `совутгич ${fmt(data.cold_storage_count,0)} та`, `${fmt(data.cold_storage_capacity_t,0)} т`];
      return ["—", "—", "—"];
    }

    function renderDistricts() {
      const c = selectedCommitment();
      const key = c.districtKey;
      $("#view-districts").innerHTML = `<div class="layout">
        <aside class="panel"><div class="panel-head"><div><h2>Мажбурият танлаш</h2><p>Туман кесими мавжуд бўлган мажбуриятлар рақамли очилади.</p></div></div>${commitmentList()}</aside>
        <section class="panel">
          <div class="panel-head"><div><h2>${esc(c.title)}</h2><p>${key ? "Туман/шаҳар кесимидаги рақамлар" : "Бу мажбурият учун туман кесими йўқ ёки алоҳида реестр талаб қилади."}</p></div></div>
          ${key ? `<div class="table-wrap"><table><thead><tr><th>Туман/шаҳар</th><th>I чорак</th><th>I ярим йиллик</th><th>Йил якуни</th><th>Бажарилмаган топшириқ</th></tr></thead>
          <tbody>${DATA.districts.map((d) => { const v = districtValue(d, key); return `<tr><td class="num">${esc(d.name)}</td><td>${esc(v[0])}</td><td>${esc(v[1])}</td><td>${esc(v[2])}</td><td>${d.debt.task_unfinished} / ${d.debt.task_total}</td></tr>`; }).join("")}</tbody></table></div>` : `<div class="empty">Туман кесими учун алоҳида манба ёки қўлда киритиладиган топшириқ реестри керак.</div>`}
        </section>
      </div>`;
      bindCommitButtons();
    }

    function renderTasks() {
      const rows = filteredCommitments().flatMap((c) => c.taskParagraphs.map((t) => ({ c, t })));
      $("#view-tasks").innerHTML = `<section class="panel">
        <div class="panel-head"><div><h2>Топшириқлар қандай келиб чиқади?</h2><p>Топшириқлар кафолат хати матнидан ажралади ва мажбурият/KPIга уланади.</p></div></div>
        <div class="table-wrap"><table><thead><tr><th>Мажбурият</th><th>Параграф</th><th>Топшириқ матни</th><th>Рақамли асос</th><th>Кейинги мониторинг</th></tr></thead>
        <tbody>${rows.map(({c,t}) => `<tr><td class="num">${esc(c.id)}<span class="sub">${esc(c.title)}</span></td><td>${t.n}-п.</td><td>${esc(t.text)}</td><td>${c.excel.map((x) => `<span class="sub">${esc(x)}</span>`).join("")}</td><td>масъул + муддат + далил киритилади</td></tr>`).join("")}</tbody></table></div>
      </section>`;
    }

    function renderSources() {
      $("#view-sources").innerHTML = `<div class="two-col">
        <section class="panel"><div class="panel-head"><div><h2>Манба файллар</h2><p>Жадваллар мажбуриятларнинг рақамли асоси.</p></div></div>
          <div class="table-wrap"><table><thead><tr><th>Файл</th><th>Варақ</th><th>Қатор</th><th>Устун</th></tr></thead><tbody>${DATA.sources.map((s) => `<tr><td class="num">${esc(s.file)}</td><td>${esc(s.sheet || "—")}</td><td>${esc(s.rows || "—")}</td><td>${esc(s.cols || "—")}</td></tr>`).join("")}</tbody></table></div>
        </section>
        <section class="panel"><div class="panel-head"><div><h2>Аудит қайдлари</h2><p>Импорт логикасида ҳисобга олинадиган муаммолар.</p></div></div>
          <div class="source-list" style="padding: 16px;">${DATA.data_quality.map((x) => `<div class="source-item"><b>${esc(x.severity)}</b> · ${esc(x.detail)}</div>`).join("")}</div>
        </section>
      </div>`;
    }

    function renderStats() {
      const full = DATA.commitments.filter((c) => c.link === "Тўлиқ").length;
      const partial = DATA.commitments.filter((c) => c.link !== "Тўлиқ").length;
      return `<div class="stats"><div class="stat"><span>Мажбуриятлар</span><strong>${DATA.commitments.length}</strong></div><div class="stat"><span>Тўлиқ боғланган</span><strong>${full}</strong></div><div class="stat"><span>Қисман боғланган</span><strong>${partial}</strong></div><div class="stat"><span>Туман/шаҳар</span><strong>${DATA.districts.length}</strong></div></div>`;
    }

    function render() {
      renderNav();
      renderFilters();
      document.querySelectorAll(".view").forEach((v) => v.classList.remove("active"));
      $(`#view-${state.view}`).classList.add("active");
      if (state.view === "map") renderMap();
      if (state.view === "matrix") renderMatrix();
      if (state.view === "districts") renderDistricts();
      if (state.view === "tasks") renderTasks();
      if (state.view === "sources") renderSources();
      const active = $(`#view-${state.view}`);
      if (!active.querySelector(".stats") && state.view !== "map") active.insertAdjacentHTML("afterbegin", renderStats());
    }

    $("#sectionFilter").addEventListener("change", (e) => { state.section = e.target.value; render(); });
    $("#search").addEventListener("input", (e) => { state.search = e.target.value; render(); });
    render();
  </script>
</body>
</html>"""


if __name__ == "__main__":
    main()
