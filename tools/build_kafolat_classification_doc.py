from __future__ import annotations

import importlib.util
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GENERATOR = ROOT / "tools" / "generate_andijon_integrated_platform_v65.py"
DATA_JSON = ROOT / "platform prototypes" / "andijon_full_pilot_assets" / "andijon_full_pilot_data.json"
OUT_DOCX = ROOT / "platform prototypes" / "andijon_kafolat_hati_topshiriqlar_klassifikatsiyasi.docx"

KPI_LABELS = {
    "grp": "ЯҲМ",
    "industry": "Саноат",
    "localization": "Маҳаллийлаштириш",
    "energy_electricity": "Электр энергияси",
    "energy_gas": "Табиий газ",
    "services": "Хизматлар",
    "agriculture": "Қишлоқ хўжалиги",
    "construction": "Қурилиш",
    "inflation": "Инфляция",
    "budget": "Бюджет",
    "budget_investment": "Бюджет инвестиция",
    "investment": "Хорижий инвестиция",
    "export": "Экспорт",
    "unemployment": "Ишсизлик",
    "jobs": "Иш ўринлари",
    "legalization": "Бандликни легаллаштириш",
    "mfy_clear": "Камбағалликдан холи маҳалла",
    "microprojects": "Микролойиҳалар",
    "poverty": "Камбағаллик",
}

TYPE_ORDER = {
    "KPI мақсад": 1,
    "Туман мақсадли кўрсаткичи": 2,
    "Ижро топшириғи": 3,
    "Маълумот/асос": 4,
}


def is_platform_kpi_target(task: dict) -> bool:
    text = str(task.get("title") or "")
    low = text.casefold()

    # Only the indicators that are represented as platform KPI cards should be
    # removed from the task register. Numeric driver targets remain tasks.
    if "ялпи ҳудудий маҳсулот" in low:
        return True
    if "саноат маҳсулот" in low and "ўсиш суръат" in low:
        return True
    if "бозор хизматлари ҳажм" in low:
        return True
    if "қишлоқ хўжалиги маҳсулотлари ҳажм" in low:
        return True
    if "қурилиш ишлари ҳажм" in low:
        return True
    if "инфляция даражас" in low:
        return True
    if "бюджет даромадлари прогнози" in low or "йиллик прогноз" in low:
        return True
    if "2-чорак учун" in low and "даромадларга эришиш" in low:
        return True
    if "инвестициялар ва кредитларни жалб қилиш" in low and "лойиҳа" not in low:
        return True
    if re.search(r"^\s*\d[\d\s,.]*\s*млн\s+доллар.*экспорт\s+таъминлаш", low):
        return True
    if "ишсизлик даражас" in low:
        return True
    if "камбағаллик даражас" in low:
        return True
    return False


def load_generator_module():
    spec = importlib.util.spec_from_file_location("andijon_v65", GENERATOR)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load generator: {GENERATOR}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def short_text(text: str, limit: int = 260) -> str:
    text = " ".join(str(text or "").split())
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


def has_target_value(text: str) -> bool:
    low = text.casefold()
    return bool(re.search(r"\d+[\d\s,.]*\s*(фоиз|%|трлн|млрд|млн|минг|та)", low))


def has_action_word(text: str) -> bool:
    low = text.casefold()
    action_words = [
        "ишга тушир",
        "ташкил эт",
        "қайта тиклаш",
        "қайта йўлга қўйиш",
        "кўмаклашиш",
        "манзилли ишлаш",
        "жойлаштир",
        "ажрат",
        "ўқит",
        "легаллаштир",
        "ўтказ",
        "ишлаб чиқ",
        "амалга ошир",
        "назорат",
        "мониторинг",
        "кредит",
        "субсид",
        "экспортчи",
        "тадбиркор",
        "фойдаланишга топшир",
    ]
    return any(word in low for word in action_words)


def classify_task(task: dict) -> tuple[str, str, str]:
    text = str(task.get("title") or "")
    low = text.casefold()
    districts = task.get("districts") or []
    group = str(task.get("group") or "")
    instrument_words = [
        "корхона",
        "лойиҳа",
        "экспортчи",
        "омбор",
        "ярмарка",
        "кўча",
        "объект",
        "хонадон",
        "уйлар",
        "субсид",
        "кредит",
        "ўқит",
        "легаллаштир",
        "ишга жойлаштир",
        "бизнес-миссия",
        "кўргазма",
        "инфратузилма",
        "хизмат кўрсатиш шаҳоб",
    ]
    has_instrument = any(word in low for word in instrument_words)
    strong_action = has_action_word(text)
    is_inflation_cap = "инфляция даражас" in low or "нархини" in low or "нархларини" in low

    if is_platform_kpi_target(task):
        return (
            "KPI мақсад",
            "KPI экрани",
            "Платформада асосий KPI сифатида киритилган индикатор. Топшириқлар реестридан чиқарилади.",
        )

    if districts and has_target_value(text):
        if has_instrument and any(word in low for word in ["ишга тушир", "ташкил эт", "корхона", "экспортчи", "лойиҳа"]):
            return (
                "Ижро топшириғи",
                "Топшириқлар / Ижро мониторинги / Туман профили",
                "Манзилли ҳудуд ва амалий ҳаракат бор. Топшириқ сифатида қолади.",
            )
        return (
            "Туман мақсадли кўрсаткичи",
            "Туманлар экрани",
            "Туман/шаҳар кесимидаги режа. Топшириқлар сонига қўшилмайди.",
        )

    if strong_action or has_target_value(text):
        return (
            "Ижро топшириғи",
            "Топшириқлар / Ижро мониторинги",
            "KPIга эришиш учун амалий ҳаракат бор. Топшириқ сифатида қолади.",
        )

    return (
        "Маълумот/асос",
        "Керак бўлса маълумот сифатида",
        "Мониторинг объекти эмас. Асословчи ёки тушунтирувчи матн сифатида қаралади.",
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table_header(table, labels: list[str]) -> None:
    hdr = table.rows[0]
    for idx, label in enumerate(labels):
        cell = hdr.cells[idx]
        set_cell_text(cell, label, bold=True, color="FFFFFF", size=8)
        set_cell_shading(cell, "1F5FBF")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.first_child_found_in("w:tcMar")
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side in ["top", "left", "bottom", "right"]:
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), "80")
                node.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = False
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = RGBColor(31, 95, 191)


def add_note_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    style_table(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF3FF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(19, 64, 133)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r2.font.size = Pt(9)


def add_summary_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    add_table_header(table, ["Кўрсаткич", "Сони", "Платформа қарори"])
    for label, count, decision in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=9)
        set_cell_text(cells[1], count, bold=True, size=9)
        set_cell_text(cells[2], decision, size=9)
    style_table(table)


def build_registry(tasks: list[dict]) -> list[dict]:
    counters = Counter()
    registry = []
    for task in tasks:
        kind, surface, note = classify_task(task)
        counters[kind] += 1
        prefix = {
            "KPI мақсад": "KPI",
            "Туман мақсадли кўрсаткичи": "D",
            "Ижро топшириғи": "T",
            "Маълумот/асос": "INFO",
        }[kind]
        new_id = f"{prefix}-{counters[kind]:03d}" if prefix not in {"KPI", "D"} else f"{prefix}-{counters[kind]:02d}"
        registry.append({
            "new_id": new_id,
            "source_id": str(task.get("id") or ""),
            "section": str(task.get("section") or ""),
            "title": str(task.get("title") or ""),
            "kpi": KPI_LABELS.get(str(task.get("kpi") or ""), str(task.get("kpi") or "")),
            "kind": kind,
            "deadline": str(task.get("deadline") or task.get("period") or ""),
            "districts": ", ".join(task.get("districts") or []) if task.get("districts") else "Вилоят",
            "surface": surface,
            "note": note,
        })
    return registry


def add_registry_table(doc: Document, title: str, rows: list[dict], limit: int | None = None) -> None:
    visible = rows[:limit] if limit else rows
    table = doc.add_table(rows=2, cols=9)
    title_cell = table.rows[0].cells[0]
    for cell in table.rows[0].cells[1:]:
        title_cell = title_cell.merge(cell)
    set_cell_text(title_cell, title, bold=True, color="1F5FBF", size=11)
    set_cell_shading(title_cell, "EAF3FF")
    headers = [
        "Янги рақам",
        "Манба ID",
        "Кафолат хати бўлими",
        "Қисқа мазмун",
        "KPI",
        "Тур",
        "Муддат",
        "Ҳудуд",
        "Платформада кўринадиган жой / изоҳ",
    ]
    hdr = table.rows[1]
    for idx, label in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_text(cell, label, bold=True, color="FFFFFF", size=8)
        set_cell_shading(cell, "1F5FBF")
    for row in visible:
        cells = table.add_row().cells
        values = [
            row["new_id"],
            row["source_id"],
            short_text(row["section"], 130),
            short_text(row["title"], 300),
            row["kpi"],
            row["kind"],
            row["deadline"],
            short_text(row["districts"], 120),
            short_text(f"{row['surface']}. {row['note']}", 240),
        ]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, bold=idx in {0, 5}, size=7)
        if row["kind"] == "KPI мақсад":
            set_cell_shading(cells[5], "EAF3FF")
        elif row["kind"] == "Туман мақсадли кўрсаткичи":
            set_cell_shading(cells[5], "FFF3CD")
        elif row["kind"] == "Ижро топшириғи":
            set_cell_shading(cells[5], "E6F4EA")
        else:
            set_cell_shading(cells[5], "F1F3F5")
    style_table(table)
    if limit and len(rows) > limit:
        p = doc.add_paragraph(f"Эслатма: ушбу бўлимда биринчи {limit} қатор кўрсатилди; тўлиқ реестр кейинги бўлимда берилган.")
        p.runs[0].font.name = "Arial"
        p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        p.runs[0].font.size = Pt(9)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def main() -> None:
    module = load_generator_module()
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    tasks = module.extract_kafolat_tasks(data)
    meta = data.get("task_meta", {})
    registry = build_registry(tasks)
    counts = Counter(row["kind"] for row in registry)
    by_kpi = defaultdict(Counter)
    for row in registry:
        by_kpi[row["kpi"]][row["kind"]] += 1

    doc = Document()
    set_document_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Андижон вилояти кафолат хати топшириқларини қайта классификация қилиш")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(19, 64, 133)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("KPI мақсадлари, туман мақсадлари ва ҳақиқий ижро топшириқларини ажратиш бўйича ишчи реестр")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(83, 100, 122)

    add_note_box(
        doc,
        "Асосий қарор",
        "Кафолат хатидаги ҳар бир рақамли натижа KPI эмас. Фақат платформада асосий KPI карточкаси сифатида киритилган индикаторлар топшириқлар сафидан чиқарилади. Давлат харидлари, маҳаллий контент, лойиҳа, корхона, энергия тежаш, ярмарка ва шунга ўхшаш KPIга эришиш драйверлари 'Ижро топшириғи' сифатида қолади.",
    )

    add_heading(doc, "1. Қисқа хулоса", 1)
    declared_total = meta.get("declared_total", 114)
    add_summary_table(
        doc,
        [
            ("Кафолат хатидаги расмий топшириқлар", str(declared_total), "Манбадаги расмий сон сифатида сақланади"),
            ("Мониторинг учун ажратилган қаторлар", str(len(registry)), "Муддат/деталлар бўйича парчаланган ишчи реестр"),
            ("KPI мақсадлар", str(counts["KPI мақсад"]), "KPI экранига ўтказилади"),
            ("Туман мақсадли кўрсаткичлари", str(counts["Туман мақсадли кўрсаткичи"]), "Туманлар экранида режа-факт-ижро сифатида юритилади"),
            ("Ҳақиқий ижро топшириқлари", str(counts["Ижро топшириғи"]), "Топшириқлар ва ижро мониторингида қолади"),
            ("Маълумот/асос", str(counts["Маълумот/асос"]), "Мониторинг топшириғи сифатида ҳисобланмайди"),
        ],
    )

    add_heading(doc, "2. Классификация қоидалари", 1)
    rules = [
        ("KPI мақсад", "Фақат платформада асосий KPI карточкаси сифатида киритилган индикатор: ЯҲМ, саноат, инфляция, бюджет, инвестиция, экспорт, ишсизлик, камбағаллик ва тасдиқланган таркибий KPI."),
        ("Туман мақсадли кўрсаткичи", "Туман/шаҳар кесимида берилган режа. У топшириқ эмас, туман мониторинги учун мақсадли кўрсаткич."),
        ("Ижро топшириғи", "KPIга эришиш учун амалий ҳаракат: лойиҳа, корхона, экспортчи, иш ўрни, кредит, субсидия, ўқитиш, легаллаштириш, мониторинг."),
        ("Маълумот/асос", "Кузатув ёки ҳисобот объекти бўлмаган тушунтирувчи матн."),
    ]
    table = doc.add_table(rows=1, cols=2)
    add_table_header(table, ["Тур", "Қоида"])
    for kind, rule in rules:
        cells = table.add_row().cells
        set_cell_text(cells[0], kind, bold=True, size=9)
        set_cell_text(cells[1], rule, size=9)
    style_table(table)

    add_heading(doc, "3. KPI кесимида тақсимот", 1)
    kpi_table = doc.add_table(rows=1, cols=5)
    add_table_header(kpi_table, ["KPI", "KPI мақсад", "Туман мақсади", "Ижро топшириғи", "Маълумот/асос"])
    for kpi in sorted(by_kpi):
        cells = kpi_table.add_row().cells
        set_cell_text(cells[0], kpi, bold=True, size=8)
        set_cell_text(cells[1], str(by_kpi[kpi]["KPI мақсад"]), size=8)
        set_cell_text(cells[2], str(by_kpi[kpi]["Туман мақсадли кўрсаткичи"]), size=8)
        set_cell_text(cells[3], str(by_kpi[kpi]["Ижро топшириғи"]), size=8)
        set_cell_text(cells[4], str(by_kpi[kpi]["Маълумот/асос"]), size=8)
    style_table(kpi_table)

    add_registry_table(doc, "4. KPI сифатида ўтказиладиган бандлар", [r for r in registry if r["kind"] == "KPI мақсад"])
    add_registry_table(doc, "5. Туман мақсадли кўрсаткичлари", [r for r in registry if r["kind"] == "Туман мақсадли кўрсаткичи"])
    add_registry_table(doc, "6. Топшириқлар экранида қоладиган ҳақиқий ижро топшириқлари", [r for r in registry if r["kind"] == "Ижро топшириғи"])
    add_registry_table(doc, "7. Тўлиқ ишчи реестр", registry)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
