from __future__ import annotations

import json
import math
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_ROOT = ROOT / "data_source"
OUT_DIR = ROOT / "platform prototypes" / "andijon_full_pilot_assets"
OUT_HTML = ROOT / "platform prototypes" / "andijon_full_pilot_prototype.html"
OUT_JSON = OUT_DIR / "andijon_full_pilot_data.json"
OUT_AUDIT = OUT_DIR / "andijon_full_pilot_audit.md"


PERIODS = {
    "q1": "I чорак",
    "h1": "I ярим йиллик",
    "m9": "9 ой",
    "year": "2026 йил",
}


DISTRICT_CANON = {
    "Андижон шаҳри": "Андижон шаҳри",
    "Андижон шаҳар": "Андижон шаҳри",
    "Андижон ш.": "Андижон шаҳри",
    "Андижон ш": "Андижон шаҳри",
    "Хонобод шаҳри": "Хонобод шаҳри",
    "Хонобод шаҳар": "Хонобод шаҳри",
    "Хонобод ш.": "Хонобод шаҳри",
    "Хонобод ш": "Хонобод шаҳри",
    "Андижон тумани": "Андижон тумани",
    "Андижон": "Андижон тумани",
    "Асака тумани": "Асака тумани",
    "Асака": "Асака тумани",
    "Балиқчи тумани": "Балиқчи тумани",
    "Балиқчи": "Балиқчи тумани",
    "Булоқбоши тумани": "Булоқбоши тумани",
    "Булоқбоши": "Булоқбоши тумани",
    "Бўстон тумани": "Бўстон тумани",
    "Бустон": "Бўстон тумани",
    "Бўстон": "Бўстон тумани",
    "Жалақудуқ тумани": "Жалақудуқ тумани",
    "Жалақудуқ": "Жалақудуқ тумани",
    "Жалолқудуқ тумани": "Жалақудуқ тумани",
    "Избоскан тумани": "Избоскан тумани",
    "Избоскан": "Избоскан тумани",
    "Улуғнор тумани": "Улуғнор тумани",
    "Улуғнор": "Улуғнор тумани",
    "Улуғноp": "Улуғнор тумани",
    "Қўрғонтепа тумани": "Қўрғонтепа тумани",
    "Қўрғонтепа": "Қўрғонтепа тумани",
    "Кўрғонтепа тумани": "Қўрғонтепа тумани",
    "Марҳамат тумани": "Марҳамат тумани",
    "Мархамат тумани": "Марҳамат тумани",
    "Мархамат": "Марҳамат тумани",
    "Олтинкўл тумани": "Олтинкўл тумани",
    "Олтинкўл": "Олтинкўл тумани",
    "Пахтаобод тумани": "Пахтаобод тумани",
    "Пахтаобод": "Пахтаобод тумани",
    "Хўжаобод тумани": "Хўжаобод тумани",
    "Хўжаобод": "Хўжаобод тумани",
    "Шаҳрихон тумани": "Шаҳрихон тумани",
    "Шахрихон тумани": "Шаҳрихон тумани",
    "Шаҳриҳон": "Шаҳрихон тумани",
    "Шахрихон": "Шаҳрихон тумани",
    "Андижон вилояти": "Андижон вилояти",
    "Анджижон вилояти": "Андижон вилояти",
    "Анджижон вилояти ": "Андижон вилояти",
    "ЖАМИ": "Андижон вилояти",
    "Жами": "Андижон вилояти",
}


DISTRICT_ORDER = [
    "Андижон шаҳри",
    "Хонобод шаҳри",
    "Андижон тумани",
    "Асака тумани",
    "Балиқчи тумани",
    "Булоқбоши тумани",
    "Бўстон тумани",
    "Жалақудуқ тумани",
    "Избоскан тумани",
    "Улуғнор тумани",
    "Қўрғонтепа тумани",
    "Марҳамат тумани",
    "Олтинкўл тумани",
    "Пахтаобод тумани",
    "Хўжаобод тумани",
    "Шаҳрихон тумани",
]


def find_andijon_folder() -> Path:
    return next(p for p in DATA_ROOT.rglob("*") if p.is_dir() and p.name.startswith("2. "))


def is_num(value) -> bool:
    return isinstance(value, (int, float)) and not isinstance(value, bool) and not math.isnan(value)


def clean_name(value) -> str:
    if value is None:
        return ""
    text = str(value).replace("\n", " ").strip()
    text = re.sub(r"\s+", " ", text)
    return text


def canon(value) -> str:
    text = clean_name(value)
    text = text.replace("p", "р")
    return DISTRICT_CANON.get(text, text)


def round_num(value, digits=1):
    if value is None or value == "":
        return None
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float)):
        return round(value, digits)
    return value


def fmt(value, digits=1, suffix=""):
    if value is None or value == "":
        return "—"
    if isinstance(value, str):
        return value
    rounded = round(value, digits)
    if rounded == int(rounded):
        text = f"{int(rounded):,}".replace(",", " ")
    else:
        text = f"{rounded:,.{digits}f}".replace(",", " ").replace(".", ",")
    return text + suffix


def status_from_ratio(actual, target, higher_is_better=True):
    if not is_num(actual) or not is_num(target) or target == 0:
        return "grey", None
    ratio = actual / target if higher_is_better else target / actual
    pct = ratio * 100
    if pct >= 100:
        status = "green"
    elif pct >= 80:
        status = "amber"
    else:
        status = "red"
    return status, round(pct, 1)


def load_wb(folder: Path, prefix: str):
    path = next(p for p in folder.glob("*.xlsx") if p.name.startswith(prefix))
    return path, openpyxl.load_workbook(path, data_only=True, read_only=True)


def source_register(folder: Path):
    records = []
    for path in sorted(folder.iterdir(), key=lambda p: p.name):
        if path.suffix.lower() == ".xlsx":
            wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
            for ws in wb.worksheets:
                nonempty = 0
                for row in ws.iter_rows(values_only=True):
                    if any(v not in (None, "") for v in row):
                        nonempty += 1
                records.append(
                    {
                        "file": path.name,
                        "type": "Excel",
                        "sheet": ws.title,
                        "rows": ws.max_row,
                        "cols": ws.max_column,
                        "nonempty_rows": nonempty,
                    }
                )
        elif path.suffix.lower() == ".docx":
            doc = Document(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            records.append(
                {
                    "file": path.name,
                    "type": "DOCX",
                    "sheet": "",
                    "rows": len(paras),
                    "cols": "",
                    "nonempty_rows": len(paras),
                }
            )
    return records


def extract_macro(folder: Path):
    path, wb = load_wb(folder, "1.1")
    regional = []
    district = defaultdict(dict)

    ws = wb["1.1. ЯҲМ"]
    for row in ws.iter_rows(min_row=6, max_row=10, values_only=True):
        indicator = clean_name(row[1])
        regional.append(
            {
                "sector": "Макро иқтисодиёт",
                "indicator": indicator,
                "unit": "млрд сўм",
                "q1_value": round_num(row[2]),
                "q1_growth": round_num(row[3]),
                "h1_value": round_num(row[4]),
                "h1_growth": round_num(row[5]),
                "m9_value": round_num(row[6]),
                "m9_growth": round_num(row[7]),
                "year_value": round_num(row[8]),
                "year_growth": round_num(row[9]),
                "source": f"{path.name} · {ws.title}",
            }
        )

    mappings = [
        ("1.2. Саноат", "industry", "Саноат маҳсулотлари", 7, 23),
        ("1.4. ҚХ", "agriculture", "Қишлоқ хўжалиги", 7, 23),
        ("1.5. Бозор хизматлари", "services", "Бозор хизматлари", 7, 23),
    ]
    for sheet, key, label, start, end in mappings:
        ws = wb[sheet]
        for row in ws.iter_rows(min_row=start, max_row=end, values_only=True):
            name = canon(row[1] if row[1] else row[0])
            if name == "Андижон вилояти" or name not in DISTRICT_ORDER:
                continue
            district[name][key] = {
                "label": label,
                "q1_value": round_num(row[2]),
                "q1_growth": round_num(row[3]),
                "h1_value": round_num(row[4]),
                "h1_growth": round_num(row[5]),
                "m9_value": round_num(row[6]),
                "m9_growth": round_num(row[7]),
                "year_value": round_num(row[8]),
                "year_growth": round_num(row[9]),
                "unit": "млрд сўм",
                "source": f"{path.name} · {sheet}",
            }
    ws = wb["1.3. Ҳудудий саноат"]
    for row in ws.iter_rows(min_row=8, max_row=23, values_only=True):
        name = canon(row[1])
        if name not in DISTRICT_ORDER:
            continue
        district[name]["localization_projects_h1"] = round_num(row[10], 0)
        district[name]["energy_electricity_h1"] = round_num(row[16])
        district[name]["energy_gas_h1"] = round_num(row[17])
    return regional, district


def extract_inflation(folder: Path):
    path, wb = load_wb(folder, "2.1")
    balance = []
    ws = wb["1.1. Баланс"]
    for row in ws.iter_rows(min_row=6, max_row=17, values_only=True):
        if not row[1]:
            continue
        balance.append(
            {
                "product": clean_name(row[1]),
                "resource_total": round_num(row[2]),
                "production": round_num(row[4]),
                "import": round_num(row[5]),
                "use_total": round_num(row[6]),
                "local_supply_ratio": round_num(row[12], 2),
                "year_end_stock": round_num(row[14]),
                "source": f"{path.name} · {ws.title}",
            }
        )
    warehouses = defaultdict(dict)
    ws = wb["1.2. Омборлар"]
    for row in ws.iter_rows(min_row=6, max_row=21, values_only=True):
        name = canon(row[1])
        if name not in DISTRICT_ORDER:
            continue
        warehouses[name] = {
            "reserve_warehouses": round_num(row[2], 0),
            "reserve_capacity_t": round_num(row[3], 0),
            "cold_storage_count": round_num(row[4], 0),
            "cold_storage_capacity_t": round_num(row[5], 0),
            "new_small_cold_storage_count": round_num(row[6], 0),
            "new_large_cold_storage_count": round_num(row[9], 0),
            "source": f"{path.name} · {ws.title}",
        }
    return balance, warehouses


def extract_budget(folder: Path):
    path, wb = load_wb(folder, "3")
    ws = wb["тушум"]
    regional = {}
    districts = defaultdict(dict)
    for row in ws.iter_rows(min_row=8, max_row=25, values_only=True):
        name = canon(row[1])
        rec = {
            "year_plan": round_num(row[2]),
            "h1_plan": round_num(row[3]),
            "q2_plan": round_num(row[4]),
            "year_expected": round_num(row[5]),
            "h1_expected": round_num(row[6]),
            "q2_expected": round_num(row[7]),
            "h1_execution_pct": round_num(row[11]),
            "q2_execution_pct": round_num(row[13]),
            "unit": "млрд сўм",
            "source": f"{path.name} · {ws.title}",
        }
        if name == "Андижон вилояти":
            regional = rec
        elif name in DISTRICT_ORDER:
            districts[name]["budget"] = rec
    return regional, districts


def extract_budget_invest(folder: Path):
    path, wb = load_wb(folder, "4.1")
    ws = wb["2.Анд"]
    regional = {}
    districts = defaultdict(dict)
    for row in ws.iter_rows(min_row=7, max_row=27, values_only=True):
        name = canon(row[1])
        rec = {
            "objects": round_num(row[2], 0),
            "limit": round_num(row[3]),
            "q1_absorption": round_num(row[4]),
            "q1_pct": round_num(row[5]),
            "h1_absorption": round_num(row[8]),
            "h1_pct": round_num(row[9]),
            "year_absorption": round_num(row[12]),
            "year_pct": round_num(row[13]),
            "commissioning_year_count": round_num(row[20], 0),
            "commissioning_year_value": round_num(row[21]),
            "unit": "млн сўм",
            "source": f"{path.name} · {ws.title}",
        }
        if name == "Андижон вилояти":
            regional = rec
        elif name in DISTRICT_ORDER:
            districts[name]["budget_investment"] = rec
    return regional, districts


def extract_foreign_invest(folder: Path):
    path, wb = load_wb(folder, "4.2")
    ws = wb["4,2-хорижий инв"]
    regional = {}
    districts = defaultdict(dict)
    for row in ws.iter_rows(min_row=7, max_row=23, values_only=True):
        name = canon(row[1] or row[2])
        rec = {
            "year_forecast": round_num(row[6]),
            "q1_plan": round_num(row[8]),
            "q1_actual": round_num(row[11]),
            "q1_pct": round_num(row[12]),
            "h1_plan": round_num(row[14]),
            "h1_expected": round_num(row[17]),
            "h1_pct": round_num(row[18]),
            "year_expected": round_num(row[20]),
            "year_pct": round_num(row[21]),
            "q1_projects": round_num(row[23], 0),
            "h1_projects": round_num(row[26], 0),
            "h1_jobs": round_num(row[28], 0),
            "unit": "млн доллар",
            "source": f"{path.name} · {ws.title}",
        }
        if name == "Андижон вилояти":
            regional = rec
        elif name in DISTRICT_ORDER:
            districts[name]["foreign_investment"] = rec
    return regional, districts


def extract_export(folder: Path):
    path, wb = load_wb(folder, "5.1")
    ws = wb["5-жадвал"]
    regional = {}
    districts = defaultdict(dict)
    for row in ws.iter_rows(min_row=6, max_row=22, values_only=True):
        name = canon(row[1] if row[1] else row[0])
        rec = {
            "year_forecast": round_num(row[2]),
            "q1_exporters": round_num(row[3], 0),
            "q1_value": round_num(row[4]),
            "q1_growth": round_num(row[5]),
            "h1_exporters": round_num(row[7], 0),
            "h1_expected": round_num(row[8]),
            "h1_growth": round_num(row[9]),
            "year_exporters": round_num(row[11], 0),
            "year_expected": round_num(row[12]),
            "year_growth": round_num(row[13]),
            "unit": "минг доллар",
            "source": f"{path.name} · {ws.title}",
        }
        if name == "Андижон вилояти":
            regional = rec
        elif name in DISTRICT_ORDER:
            districts[name]["export"] = rec
    return regional, districts


def extract_employment(folder: Path):
    path, wb = load_wb(folder, "6")
    ws = wb["6. Камбағаллик"]
    regional = {}
    districts = defaultdict(dict)
    for row in ws.iter_rows(min_row=7, max_row=23, values_only=True):
        name = canon(row[1] if row[1] else row[0])
        rec = {
            "unemployment_h1": round_num(row[2], 2),
            "unemployment_year": round_num(row[3], 2),
            "poverty_h1": round_num(row[4], 2),
            "poverty_year": round_num(row[5], 2) if is_num(row[5]) else row[5],
            "mfy_h1": round_num(row[6], 0),
            "mfy_year": round_num(row[7], 0),
            "jobs_h1": round_num(row[8], 3),
            "jobs_year": round_num(row[9], 3),
            "legalization_h1": round_num(row[10], 3),
            "legalization_year": round_num(row[11], 3),
            "microprojects_h1": round_num(row[12], 0),
            "microprojects_year": round_num(row[13], 0),
            "unit": "минг нафар / %",
            "source": f"{path.name} · {ws.title}",
        }
        if name == "Андижон вилояти":
            regional = rec
        elif name in DISTRICT_ORDER:
            districts[name]["employment"] = rec
    return regional, districts


def extract_docx_tasks(folder: Path):
    path = next(folder.glob("*.docx"))
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tasks = []
    sector = "Умумий"
    sector_patterns = [
        ("макро", "Макро иқтисодиёт"),
        ("Саноат", "Макро иқтисодиёт"),
        ("Қишлоқ", "Қишлоқ хўжалиги"),
        ("Бюджет", "Бюджет"),
        ("Инвестиция", "Хорижий инвестиция"),
        ("экспорт", "Экспорт"),
        ("Камбағал", "Бандлик ва камбағаллик"),
        ("ишсиз", "Бандлик ва камбағаллик"),
        ("банд", "Бандлик ва камбағаллик"),
    ]
    task_id = 1
    for i, text in enumerate(paras, start=1):
        clean = re.sub(r"\s+", " ", text)
        for pattern, mapped in sector_patterns:
            if pattern.lower() in clean.lower() and len(clean) < 140:
                sector = mapped
        has_number = bool(re.search(r"\d", clean))
        task_like = (
            clean.startswith(("–", "-", "•"))
            or re.match(r"^\d+(\.\d+)*\.", clean)
            or any(k in clean.lower() for k in ["таъминланади", "эришилади", "ишга тушир", "жойлаштир", "кўмаклаш", "тиклаш", "қисқартир"])
        )
        if task_like and has_number and 35 <= len(clean) <= 520:
            tasks.append(
                {
                    "id": f"GL-{task_id:03d}",
                    "sector": sector,
                    "source": f"{path.name} · paragraph {i}",
                    "title": clean.lstrip("–-• ").strip(),
                    "status": "grey",
                    "owner": "Вилоят/туман ҳокимияти",
                    "period": "I ярим йиллик / 2026 йил",
                }
            )
            task_id += 1
    return tasks, paras


def make_monitoring_rows(regional, districts):
    rows = []

    def add(row):
        row["id"] = f"M-{len(rows)+1:03d}"
        rows.append(row)

    for rec in regional.get("macro", []):
        add(
            {
                "scope": "region",
                "district": "Андижон вилояти",
                "sector": rec["sector"],
                "indicator": rec["indicator"],
                "period": "I ярим йиллик",
                "target": rec.get("h1_value"),
                "actual": rec.get("q1_value"),
                "unit": rec.get("unit"),
                "status": "grey",
                "execution_pct": None,
                "evidence": "I чорак факт бор, H1 мақсад билан кузатилади",
                "source": rec["source"],
                "direction": "higher",
            }
        )

    def add_district_indicator(d, sector, indicator, period, target, actual, unit, source, direction="higher"):
        status, pct = status_from_ratio(actual, target, direction == "higher")
        add(
            {
                "scope": "district",
                "district": d,
                "sector": sector,
                "indicator": indicator,
                "period": period,
                "target": target,
                "actual": actual,
                "unit": unit,
                "status": status,
                "execution_pct": pct,
                "evidence": "Excel манбасидан автоматик импорт",
                "source": source,
                "direction": direction,
            }
        )

    for d, data in districts.items():
        if "industry" in data:
            r = data["industry"]
            add_district_indicator(d, "Макро иқтисодиёт", "Саноат маҳсулотлари", "I ярим йиллик", r["h1_value"], r["q1_value"], "млрд сўм", r["source"])
        if "agriculture" in data:
            r = data["agriculture"]
            add_district_indicator(d, "Қишлоқ хўжалиги", "Қишлоқ хўжалиги маҳсулотлари", "I ярим йиллик", r["h1_value"], r["q1_value"], "млрд сўм", r["source"])
        if "services" in data:
            r = data["services"]
            add_district_indicator(d, "Хизматлар", "Бозор хизматлари", "I ярим йиллик", r["h1_value"], r["q1_value"], "млрд сўм", r["source"])
        if "budget" in data:
            r = data["budget"]
            add_district_indicator(d, "Бюджет", "Бюджет тушумлари", "I ярим йиллик", r["h1_plan"], r["h1_expected"], "млрд сўм", r["source"])
        if "budget_investment" in data:
            r = data["budget_investment"]
            add_district_indicator(d, "Бюджет инвестициялари", "Ўзлаштириш", "I ярим йиллик", r["limit"], r["h1_absorption"], "млн сўм", r["source"])
        if "foreign_investment" in data:
            r = data["foreign_investment"]
            add_district_indicator(d, "Хорижий инвестиция", "Хорижий инвестициялар", "I чорак", r["q1_plan"], r["q1_actual"], "млн доллар", r["source"])
        if "export" in data:
            r = data["export"]
            add_district_indicator(d, "Экспорт", "Экспорт ҳажми", "I ярим йиллик", r["h1_expected"], r["q1_value"], "минг доллар", r["source"])
        if "employment" in data:
            r = data["employment"]
            add_district_indicator(d, "Бандлик ва камбағаллик", "Ишсизлик даражаси", "I ярим йиллик", r["unemployment_h1"], r["unemployment_h1"], "%", r["source"], "lower")
            poverty_actual = r["poverty_h1"] if is_num(r["poverty_h1"]) else None
            add_district_indicator(d, "Бандлик ва камбағаллик", "Камбағаллик даражаси", "I ярим йиллик", r["poverty_h1"], poverty_actual, "%", r["source"], "lower")
    return rows


def compute_task_debt(districts, rows, tasks):
    by_district = defaultdict(lambda: {"red": 0, "amber": 0, "grey": 0, "green": 0})
    for r in rows:
        by_district[r["district"]][r["status"]] += 1
    result = {}
    for d in DISTRICT_ORDER:
        counts = by_district[d]
        undone = counts["red"] + counts["amber"] + counts["grey"]
        total = sum(counts.values())
        result[d] = {
            "monitoring_total": total,
            "attention": undone,
            "green": counts["green"],
            "red": counts["red"],
            "amber": counts["amber"],
            "grey": counts["grey"],
            "task_total": total,
            "task_unfinished": undone,
        }
    result["Андижон вилояти"] = {
        "monitoring_total": len([r for r in rows if r["district"] == "Андижон вилояти"]),
        "attention": len([r for r in rows if r["district"] == "Андижон вилояти" and r["status"] != "green"]),
        "task_total": len(tasks),
        "task_unfinished": len(tasks),
    }
    return result


def build_dataset():
    folder = find_andijon_folder()
    source_records = source_register(folder)
    regional = {}
    districts = defaultdict(dict)

    macro_regional, macro_district = extract_macro(folder)
    regional["macro"] = macro_regional
    for d, data in macro_district.items():
        districts[d].update(data)

    balance, warehouses = extract_inflation(folder)
    regional["food_balance"] = balance
    for d, data in warehouses.items():
        districts[d]["warehouses"] = data

    budget_regional, budget_district = extract_budget(folder)
    regional["budget"] = budget_regional
    for d, data in budget_district.items():
        districts[d].update(data)

    budget_inv_regional, budget_inv_district = extract_budget_invest(folder)
    regional["budget_investment"] = budget_inv_regional
    for d, data in budget_inv_district.items():
        districts[d].update(data)

    inv_regional, inv_district = extract_foreign_invest(folder)
    regional["foreign_investment"] = inv_regional
    for d, data in inv_district.items():
        districts[d].update(data)

    export_regional, export_district = extract_export(folder)
    regional["export"] = export_regional
    for d, data in export_district.items():
        districts[d].update(data)

    emp_regional, emp_district = extract_employment(folder)
    regional["employment"] = emp_regional
    for d, data in emp_district.items():
        districts[d].update(data)

    tasks, guarantee_paragraphs = extract_docx_tasks(folder)
    monitoring_rows = make_monitoring_rows(regional, districts)
    task_debt = compute_task_debt(districts, monitoring_rows, tasks)

    district_records = []
    for d in DISTRICT_ORDER:
        district_records.append(
            {
                "name": d,
                "owner": "туман ҳокимияти" if "тумани" in d else "шаҳар ҳокимияти",
                "data": districts.get(d, {}),
                "debt": task_debt.get(d, {}),
            }
        )

    data_quality = [
        {"issue": "region_name_typo", "detail": "Айрим варақларда 'Анджижон вилояти' ёзилган", "severity": "medium"},
        {"issue": "district_variants", "detail": "Марҳамат/Мархамат, Шаҳрихон/Шахрихон/Шаҳриҳон, Бўстон/Бустон вариантлари нормаллаштирилди", "severity": "medium"},
        {"issue": "latin_p_in_cyrillic", "detail": "6-жадвалда 'Улуғноp' сўзида Latin p бор", "severity": "low"},
        {"issue": "formula_error", "detail": "4.2-жадвал захира қаторларида #DIV/0! учрайди", "severity": "medium"},
        {"issue": "merged_headers", "detail": "Excel сарлавҳалари кўп қаторли ва merged; импорт sheet-specific parser талаб қилади", "severity": "high"},
    ]

    dataset = {
        "meta": {
            "region": "Андижон вилояти",
            "pilot": True,
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "source_folder": str(folder.relative_to(ROOT)),
            "source_file_count": len(list(folder.iterdir())),
        },
        "sources": source_records,
        "regional": regional,
        "districts": district_records,
        "monitoring": monitoring_rows,
        "tasks": tasks,
        "guarantee_letter_paragraph_count": len(guarantee_paragraphs),
        "data_quality": data_quality,
    }
    return dataset


def write_audit(data):
    red = len([r for r in data["monitoring"] if r["status"] == "red"])
    amber = len([r for r in data["monitoring"] if r["status"] == "amber"])
    green = len([r for r in data["monitoring"] if r["status"] == "green"])
    grey = len([r for r in data["monitoring"] if r["status"] == "grey"])
    lines = [
        "# Andijan Full Pilot Extraction Audit",
        "",
        f"Generated: `{data['meta']['generated_at']}`",
        f"Source folder: `{data['meta']['source_folder']}`",
        "",
        "## Coverage",
        "",
        f"- Source files/sheets registered: `{len(data['sources'])}` source records from `{data['meta']['source_file_count']}` files.",
        f"- District/city profiles: `{len(data['districts'])}`.",
        f"- Monitoring rows generated from Excel tables: `{len(data['monitoring'])}`.",
        f"- Guarantee-letter task candidates extracted: `{len(data['tasks'])}`.",
        f"- Status distribution: green `{green}`, amber `{amber}`, red `{red}`, grey `{grey}`.",
        "",
        "## Data Quality Notes",
        "",
    ]
    for item in data["data_quality"]:
        lines.append(f"- **{item['severity']}** · `{item['issue']}`: {item['detail']}")
    lines += [
        "",
        "## Main Extracted Regional KPIs",
        "",
        "| KPI | H1 | Year | Source |",
        "| --- | ---: | ---: | --- |",
    ]
    for r in data["regional"]["macro"]:
        lines.append(f"| {r['indicator']} | {fmt(r['h1_value'])} | {fmt(r['year_value'])} | {r['source']} |")
    inv = data["regional"]["foreign_investment"]
    exp = data["regional"]["export"]
    emp = data["regional"]["employment"]
    bud = data["regional"]["budget"]
    lines += [
        f"| Хорижий инвестиция | {fmt(inv.get('h1_expected'))} млн $ | {fmt(inv.get('year_expected'))} млн $ | {inv.get('source')} |",
        f"| Экспорт | {fmt(exp.get('h1_expected'))} минг $ | {fmt(exp.get('year_expected'))} минг $ | {exp.get('source')} |",
        f"| Бюджет тушумлари | {fmt(bud.get('h1_expected'))} млрд сўм | {fmt(bud.get('year_expected'))} млрд сўм | {bud.get('source')} |",
        f"| Ишсизлик | {fmt(emp.get('unemployment_h1'),2,'%')} | {fmt(emp.get('unemployment_year'),2,'%')} | {emp.get('source')} |",
        f"| Камбағаллик | {fmt(emp.get('poverty_h1'),2,'%')} | {fmt(emp.get('poverty_year'),2,'%')} | {emp.get('source')} |",
    ]
    OUT_AUDIT.write_text("\n".join(lines), encoding="utf-8")


def html_template(data_json: str) -> str:
    return f"""<!doctype html>
<html lang="uz-Cyrl">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Андижон вилояти — full pilot monitoring</title>
  <style>
    :root {{
      --ink:#0b2137; --muted:#5d7088; --paper:#eef4fa; --surface:#fbfdff; --line:#d4dfeb;
      --blue:#145fb8; --blue-soft:#e8f2ff; --green:#087b53; --green-soft:#e3f6ee;
      --amber:#a76300; --amber-soft:#fff0d0; --red:#bd3126; --red-soft:#ffe3df; --nav:#061a2d;
    }}
    * {{ box-sizing:border-box; }}
    body {{
      margin:0; color:var(--ink); font-family:"Segoe UI",Arial,sans-serif; line-height:1.35;
      background:linear-gradient(rgba(183,204,225,.5) 1px,transparent 1px),linear-gradient(90deg,rgba(183,204,225,.5) 1px,transparent 1px),var(--paper);
      background-size:34px 34px;
    }}
    button,input,select {{ font:inherit; }} button {{ cursor:pointer; }}
    :focus-visible {{ outline:3px solid color-mix(in oklch,var(--blue) 70%,white); outline-offset:3px; }}
    .top {{
      color:#f7fbff; background:linear-gradient(135deg,#234d6b,#638aa2); position:sticky; top:0; z-index:10;
      box-shadow:0 12px 30px rgba(7,26,45,.2);
    }}
    .mast {{ display:grid; grid-template-columns:minmax(0,1fr) auto minmax(0,1fr); gap:24px; align-items:center; padding:18px clamp(20px,3vw,42px) 12px; }}
    .brand {{ display:flex; gap:16px; align-items:center; min-width:0; }}
    .brand > div {{ min-width:0; }}
    .logo {{ width:110px; height:54px; object-fit:contain; filter:brightness(0) invert(1); }}
    .brand h1 {{ margin:0; font-size:clamp(20px,2.2vw,28px); line-height:1.1; }}
    .brand p {{ margin:5px 0 0; color:rgba(247,251,255,.78); font-size:13px; font-weight:650; overflow-wrap:anywhere; }}
    .kmark {{ font-size:clamp(42px,5vw,64px); letter-spacing:8px; font-weight:900; color:rgba(247,251,255,.72); }}
    .year {{ justify-self:end; text-align:right; font-weight:900; }} .year strong {{ display:block; font-size:46px; line-height:.95; }} .year span {{ font-size:25px; color:rgba(247,251,255,.76); }}
    .kpis {{ display:grid; grid-template-columns:repeat(6,minmax(0,1fr)); background:rgba(255,255,255,.98); color:var(--ink); border-top:1px solid rgba(255,255,255,.35); }}
    .kpi {{ border:0; border-right:1px solid var(--line); background:transparent; padding:17px clamp(14px,1.8vw,24px); text-align:left; min-height:104px; }}
    .kpi:hover,.kpi.active {{ background:var(--blue-soft); }} .kpi b {{ display:block; color:#5b6878; text-transform:uppercase; font-size:14px; }} .kpi strong {{ display:block; color:var(--blue); font-size:31px; line-height:1; margin-top:4px; }} .kpi span {{ display:block; color:#5b6b7f; font-size:13px; font-weight:800; margin-top:4px; }}
    .shell {{ display:grid; grid-template-columns:224px minmax(0,1fr); }}
    .nav {{ background:var(--nav); color:#d6e7f8; min-height:calc(100vh - 218px); padding:18px 14px; position:sticky; top:218px; height:calc(100vh - 218px); overflow:auto; }}
    .nav-title {{ padding:8px 10px 16px; border-bottom:1px solid rgba(255,255,255,.13); margin-bottom:12px; }} .nav-title b {{ color:#fff; display:block; }} .nav-title span {{ color:#9db6cf; font-size:12px; }}
    .nav button {{ width:100%; display:grid; grid-template-columns:30px 1fr auto; gap:8px; align-items:center; border:1px solid transparent; border-radius:10px; padding:10px; background:transparent; color:inherit; text-align:left; margin-bottom:6px; }}
    .nav button.active {{ background:#eef6ff; color:var(--nav); border-color:#fff; }} .num {{ width:27px; height:27px; border:1px solid currentColor; border-radius:999px; display:grid; place-items:center; font-size:12px; }}
    main {{ padding:26px clamp(20px,3vw,36px) 44px; min-width:0; }}
    .mobile-nav {{ display:none; gap:6px; overflow:auto; padding-bottom:14px; }} .mobile-nav button {{ white-space:nowrap; border:1px solid var(--line); border-radius:999px; background:#fff; padding:8px 12px; font-weight:800; }} .mobile-nav button.active {{ color:#fff; background:var(--blue); }}
    .head {{ display:grid; grid-template-columns:minmax(0,1fr) auto; gap:18px; align-items:end; margin-bottom:18px; }}
    .head h2 {{ margin:0; font-size:clamp(28px,3.3vw,42px); line-height:1.03; }} .head p {{ margin:8px 0 0; color:var(--muted); max-width:820px; }}
    .filters {{ display:flex; gap:8px; flex-wrap:wrap; justify-content:flex-end; }} .control {{ height:38px; border:1px solid #b9cadd; border-radius:9px; background:#fff; padding:0 12px; min-width:150px; }}
    .cmd {{ height:38px; border:1px solid var(--blue); border-radius:9px; background:var(--blue); color:#fff; padding:0 13px; font-weight:850; }} .cmd.secondary {{ background:#fff; color:var(--blue); }}
    .view {{ display:none; }} .view.active {{ display:block; }}
    .surface {{ border:1px solid var(--line); border-radius:10px; background:rgba(255,255,255,.91); overflow:hidden; box-shadow:0 12px 30px rgba(24,61,92,.06); }}
    .surface-head {{ display:flex; gap:12px; justify-content:space-between; align-items:center; padding:15px 17px; border-bottom:1px solid var(--line); }} .surface-head h3 {{ margin:0; font-size:17px; }} .surface-head p {{ margin:4px 0 0; color:var(--muted); font-size:13px; }}
    .scoreline {{ display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); border:1px solid var(--line); border-radius:10px; background:rgba(255,255,255,.92); overflow:hidden; margin-bottom:16px; }}
    .score {{ padding:15px 17px; border-right:1px solid var(--line); }} .score span {{ color:var(--muted); font-size:12px; font-weight:850; }} .score strong {{ display:block; font-size:32px; line-height:1; margin-top:5px; }}
    .grid2 {{ display:grid; grid-template-columns:minmax(0,1.2fr) minmax(320px,.8fr); gap:16px; align-items:start; }} .grid3 {{ display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:10px; padding:10px; }} .grid4 {{ display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:10px; padding:10px; }}
    .chip {{ display:inline-flex; border-radius:999px; padding:5px 9px; font-size:12px; font-weight:900; white-space:nowrap; }} .green {{ color:var(--green); background:var(--green-soft); border:1px solid #a3dfc5; }} .amber {{ color:var(--amber); background:var(--amber-soft); border:1px solid #ffd28a; }} .red {{ color:var(--red); background:var(--red-soft); border:1px solid #ffb5ad; }} .grey {{ color:#667789; background:#edf2f7; border:1px solid #d4dee9; }} .blue {{ color:var(--blue); background:var(--blue-soft); border:1px solid #bad8ff; }}
    .rowbtn,.cardbtn {{ width:100%; border:1px solid var(--line); border-radius:9px; background:#fbfdff; color:inherit; padding:11px; text-align:left; }} .rowbtn:hover,.cardbtn:hover,.cardbtn.active {{ border-color:var(--blue); background:var(--blue-soft); }}
    .rowbtn {{ display:grid; grid-template-columns:34px minmax(0,1fr) auto; gap:10px; align-items:center; margin:8px 0; }} .rowbtn h4,.cardbtn h4 {{ margin:0; font-size:14px; line-height:1.25; }} .rowbtn p,.cardbtn p {{ margin:5px 0 0; color:var(--muted); font-size:12px; }}
    .pad {{ padding:10px; }} .empty {{ padding:12px; color:var(--muted); }}
    .workspace {{ display:grid; grid-template-columns:300px minmax(0,1fr); gap:16px; align-items:start; }} .list {{ padding:8px; display:grid; gap:6px; }}
    .profile {{ display:grid; gap:14px; }} .profile-head {{ border:1px solid var(--line); border-radius:10px; background:rgba(255,255,255,.92); padding:17px; display:grid; grid-template-columns:1fr auto; gap:16px; }}
    .profile-head h3 {{ margin:0; font-size:25px; }} .profile-head p {{ color:var(--muted); margin:6px 0 0; }} .debt {{ border-left:1px solid var(--line); padding-left:18px; min-width:160px; }} .debt strong {{ color:var(--red); font-size:34px; display:block; line-height:1; }}
    .mini-kpis {{ display:grid; grid-template-columns:repeat(6,minmax(0,1fr)); border:1px solid var(--line); border-radius:10px; overflow:hidden; background:#fff; }} .mini {{ padding:13px; border-right:1px solid var(--line); }} .mini b {{ color:#5b6878; font-size:12px; text-transform:uppercase; }} .mini strong {{ color:var(--blue); font-size:23px; display:block; margin-top:6px; }} .mini span {{ color:var(--muted); font-size:12px; }}
    .table-wrap {{ overflow:auto; }} table {{ width:100%; border-collapse:collapse; min-width:900px; }} th,td {{ padding:12px 13px; border-bottom:1px solid var(--line); text-align:left; vertical-align:top; font-size:13px; }} th {{ background:#f0f6fb; color:#536476; text-transform:uppercase; font-size:12px; }} tr.action:hover {{ background:#f8fbff; cursor:pointer; }}
    .drawer-mask {{ position:fixed; inset:0; background:rgba(8,23,38,.34); opacity:0; pointer-events:none; z-index:29; }} .drawer-mask.open {{ opacity:1; pointer-events:auto; }} .drawer {{ position:fixed; inset:0 0 0 auto; width:min(470px,100vw); background:#fbfdff; transform:translateX(104%); transition:transform .2s ease; z-index:30; display:grid; grid-template-rows:auto 1fr; box-shadow:-22px 0 50px rgba(8,23,38,.24); }} .drawer.open {{ transform:translateX(0); }} .drawer-head {{ padding:18px; border-bottom:1px solid var(--line); display:flex; justify-content:space-between; gap:12px; }} .drawer-head h2 {{ margin:0; font-size:20px; }} .drawer-head p {{ margin:6px 0 0; color:var(--muted); font-size:13px; }} .x {{ width:36px; height:36px; border:1px solid var(--line); border-radius:9px; background:#fff; }} .drawer-body {{ overflow:auto; padding:16px; display:grid; gap:12px; }} .drawer-card {{ border:1px solid var(--line); border-radius:10px; padding:14px; background:#fff; }} .drawer-card h3 {{ margin:0 0 8px; font-size:14px; }} .drawer-card p {{ margin:0; color:var(--muted); font-size:13px; }} .kv {{ display:grid; grid-template-columns:150px 1fr; gap:7px; margin:0; }} .kv dt {{ color:var(--muted); font-size:12px; }} .kv dd {{ margin:0; font-weight:800; font-size:13px; overflow-wrap:anywhere; }}
    .mt {{ margin-top:16px; }} .text-green {{ color:var(--green); }} .text-red {{ color:var(--red); }} .text-amber {{ color:var(--amber); }}
    @media (max-width:1220px) {{ .kpis {{ grid-template-columns:repeat(3,1fr); }} .shell {{ grid-template-columns:1fr; }} .nav {{ display:none; }} .top {{ position:relative; }} .mobile-nav {{ display:flex; }} main {{ padding-top:14px; }} }}
    @media (max-width:900px) {{ .mast,.head,.grid2,.workspace {{ grid-template-columns:1fr; }} .kmark {{ display:none; }} .year {{ justify-self:start; text-align:left; }} .filters {{ justify-content:flex-start; }} .scoreline,.grid3,.grid4,.mini-kpis {{ grid-template-columns:repeat(2,1fr); }} .profile-head {{ grid-template-columns:1fr; }} .debt {{ border-left:0; border-top:1px solid var(--line); padding-left:0; padding-top:14px; }} }}
    @media (max-width:620px) {{ .mast {{ padding:18px 20px 12px; }} .brand {{ display:grid; grid-template-columns:82px minmax(0,1fr); align-items:flex-start; gap:12px; }} .logo {{ width:82px; height:44px; }} .brand h1 {{ font-size:18px; }} .brand p {{ font-size:12px; line-height:1.35; }} .kpis,.scoreline,.grid3,.grid4,.mini-kpis {{ grid-template-columns:1fr; }} .kpi {{ border-right:0; border-bottom:1px solid var(--line); }} .rowbtn {{ grid-template-columns:32px 1fr; }} .rowbtn .chip {{ grid-column:2; justify-self:start; }} }}
  </style>
</head>
<body>
  <header class="top">
    <div class="mast">
      <div class="brand"><img class="logo" src="claude/v2/images/cerr_logo-removebg-preview.png" alt="CERR"><div><h1>Андижон вилояти</h1><p>тўлиқ pilot · KPI · топшириқ · далил</p></div></div>
      <div class="kmark">KPI</div>
      <div class="year"><strong>2026</strong><span>I ярим йиллик</span></div>
    </div>
    <div class="kpis" id="kpiRail"></div>
  </header>
  <div class="shell">
    <aside class="nav"><div class="nav-title"><b>Ҳудудлар мониторинги</b><span>Андижон тўлиқ pilot</span></div><nav id="nav"></nav></aside>
    <main>
      <div class="mobile-nav" id="mobileNav"></div>
      <div class="head"><div><h2 id="pageTitle">Бошқарув маркази</h2><p id="pageSub">Андижон папкасидаги 8 та манба асосида тузилган тўлиқ pilot dataset.</p></div><div class="filters"><select class="control" id="sectorFilter"></select><select class="control" id="statusFilter"><option value="">Барча ҳолатлар</option><option value="red">Қизил</option><option value="amber">Сариқ</option><option value="green">Яшил</option><option value="grey">Кулранг</option></select><input class="control" id="search" placeholder="туман, KPI, манба..."><button class="cmd secondary" id="printBtn">Чоп этиш</button><button class="cmd" id="exportBtn">Экспорт</button></div></div>
      <section class="view active" id="view-dashboard"></section>
      <section class="view" id="view-districts"></section>
      <section class="view" id="view-monitoring"></section>
      <section class="view" id="view-tasks"></section>
      <section class="view" id="view-sources"></section>
      <section class="view" id="view-quality"></section>
    </main>
  </div>
  <div class="drawer-mask" id="mask"></div><aside class="drawer" id="drawer" role="dialog" aria-modal="true" aria-labelledby="drawerTitle" aria-hidden="true"><div class="drawer-head"><div><h2 id="drawerTitle">Тафсилот</h2><p id="drawerSub">манба ва ҳисоб</p></div><button class="x" id="closeDrawer" aria-label="Ёпиш">×</button></div><div class="drawer-body" id="drawerBody"></div></aside>
  <script>
    const DATA = {data_json};
    const state = {{ view:'dashboard', district:'Андижон шаҳри', sector:'', status:'', query:'' }};
    const $ = s => document.querySelector(s); const $$ = s => Array.from(document.querySelectorAll(s));
    const esc = v => String(v ?? '').replace(/[&<>"']/g, ch => ({{'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#039;'}}[ch]));
    const statusClass = s => s === 'green' ? 'green' : s === 'amber' ? 'amber' : s === 'red' ? 'red' : 'grey';
    const statusText = s => ({{green:'Яшил', amber:'Сариқ', red:'Қизил', grey:'Кулранг'}}[s] || s);
    const fmt = (v, suffix='') => v === null || v === undefined || v === '' ? '—' : (typeof v === 'number' ? (Math.round(v*10)/10).toLocaleString('ru-RU') : v) + suffix;
    const navItems = [['dashboard','Бошқарув','01'],['districts','Туманлар','02'],['monitoring','Мониторинг','03'],['tasks','Топшириқлар','04'],['sources','Манбалар','05'],['quality','Сифат','06']];
    function filteredRows() {{
      return DATA.monitoring.filter(r => (!state.sector || r.sector === state.sector) && (!state.status || r.status === state.status) && (!state.query || `${{r.id}} ${{r.district}} ${{r.sector}} ${{r.indicator}} ${{r.source}}`.toLowerCase().includes(state.query)));
    }}
    function init() {{
      $('#sectorFilter').innerHTML = '<option value="">Барча йўналишлар</option>' + [...new Set(DATA.monitoring.map(r=>r.sector))].sort().map(s=>`<option>${{esc(s)}}</option>`).join('');
      $('#sectorFilter').addEventListener('change', e=>{{state.sector=e.target.value; render();}});
      $('#statusFilter').addEventListener('change', e=>{{state.status=e.target.value; render();}});
      $('#search').addEventListener('input', e=>{{state.query=e.target.value.trim().toLowerCase(); render();}});
      $('#printBtn').addEventListener('click', ()=>window.print());
      $('#exportBtn').addEventListener('click', exportRows);
      $('#closeDrawer').addEventListener('click', closeDrawer); $('#mask').addEventListener('click', closeDrawer); document.addEventListener('keydown', e=>{{ if(e.key==='Escape') closeDrawer(); }});
      render();
    }}
    function render() {{
      renderNav(); renderKpis(); $$('.view').forEach(v=>v.classList.toggle('active', v.id === `view-${{state.view}}`));
      const titles = {{dashboard:['Бошқарув маркази','Андижон папкасидаги барча манбалардан тузилган pilot мониторинг.'],districts:['Туманлар профили','Ҳар бир туман/шаҳар кесимида KPI, топшириқ қарзи ва манба.'],monitoring:['Мониторинг реестри','Excel жадваллардан тузилган мақсад-факт-назорат қаторлари.'],tasks:['Кафолат хати топшириқлари','DOCX матнидан ажратилган task candidate қаторлари.'],sources:['Манбалар реестри','Файл, варақ, қаторлар ва импорт қамрови.'],quality:['Маълумот сифати','Нормаллаштириш ва импортда эҳтиёт бўладиган нуқталар.']}};
      $('#pageTitle').textContent = titles[state.view][0]; $('#pageSub').textContent = titles[state.view][1];
      if(state.view==='dashboard') renderDashboard(); if(state.view==='districts') renderDistricts(); if(state.view==='monitoring') renderMonitoring(); if(state.view==='tasks') renderTasks(); if(state.view==='sources') renderSources(); if(state.view==='quality') renderQuality();
    }}
    function renderNav() {{
      const counts = {{dashboard:'pilot', districts:DATA.districts.length, monitoring:filteredRows().length, tasks:DATA.tasks.length, sources:DATA.sources.length, quality:DATA.data_quality.length}};
      const html = navItems.map(([id,label,no])=>`<button class="${{state.view===id?'active':''}}" data-view="${{id}}" ${{state.view===id?'aria-current="page"':''}}><span class="num">${{no}}</span><b>${{label}}</b><small>${{counts[id]}}</small></button>`).join('');
      $('#nav').innerHTML=html; $('#mobileNav').innerHTML=navItems.map(([id,label])=>`<button class="${{state.view===id?'active':''}}" data-view="${{id}}">${{label}}</button>`).join('');
      $$('[data-view]').forEach(b=>b.addEventListener('click',()=>{{state.view=b.dataset.view; render();}}));
    }}
    function renderKpis() {{
      const macro = DATA.regional.macro; const inv=DATA.regional.foreign_investment; const exp=DATA.regional.export; const bud=DATA.regional.budget; const emp=DATA.regional.employment;
      const kpis = [
        ['ЯҲМ','+'+fmt((macro[0]?.h1_growth||0)-100,1)+'%','H1 · '+fmt(macro[0]?.h1_value)+' млрд сўм'],
        ['Инфляция','2,9%','II чорак лимити · кафолат хати'],
        ['Бюджет','+'+fmt((bud.h1_execution_pct||100)-100,1)+'%','H1 кутилиш '+fmt(bud.h1_expected)+' млрд'],
        ['Инвестиция',fmt(inv.h1_expected)+' млн $','H1 · '+fmt(inv.h1_pct*100,'%')],
        ['Экспорт',fmt(exp.h1_expected)+' минг $','H1 · '+fmt(exp.h1_growth,'%')],
        ['Ишсизлик',fmt(emp.unemployment_h1, '%'),'H1 паст яхши']
      ];
      $('#kpiRail').innerHTML = kpis.map(k=>`<button class="kpi"><b>${{esc(k[0])}}</b><strong>${{esc(k[1])}}</strong><span>${{esc(k[2])}}</span></button>`).join('');
    }}
    function renderDashboard() {{
      const rows=filteredRows(); const red=rows.filter(r=>r.status==='red').length; const amber=rows.filter(r=>r.status==='amber').length; const green=rows.filter(r=>r.status==='green').length; const grey=rows.filter(r=>r.status==='grey').length;
      const attention=rows.filter(r=>r.status==='red'||r.status==='amber').slice(0,10);
      const topDistricts=[...DATA.districts].sort((a,b)=>(b.debt?.attention||0)-(a.debt?.attention||0)).slice(0,8);
      $('#view-dashboard').innerHTML = `<div class="scoreline"><div class="score"><span>Мониторинг қаторлари</span><strong>${{rows.length}}</strong></div><div class="score"><span>Яшил</span><strong class="text-green">${{green}}</strong></div><div class="score"><span>Қизил/сариқ</span><strong class="text-red">${{red+amber}}</strong></div><div class="score"><span>Кулранг</span><strong class="text-amber">${{grey}}</strong></div></div><div class="grid2"><div class="surface"><div class="surface-head"><div><h3>Биринчи навбатда кўриладиган қаторлар</h3><p>Қаторни очганда манба файл, варақ ва ҳисоб ҳолати кўринади.</p></div><span class="chip red">${{attention.length}} қатор</span></div><div class="pad">${{attention.map((r,i)=>attentionRow(r,i)).join('') || '<p class="empty">Фильтр бўйича хавф йўқ.</p>'}}</div></div><div class="surface"><div class="surface-head"><div><h3>Туманлар топшириқ қарзи</h3><p>Рейтинг эмас; бажарилмаган / жами мониторинг қаторлари.</p></div></div><div class="grid3">${{topDistricts.map(d=>districtCard(d)).join('')}}</div></div></div>`;
      bindRows(); bindDistricts();
    }}
    function attentionRow(r,i) {{ return `<button class="rowbtn" data-row="${{r.id}}"><span class="num">${{i+1}}</span><span><h4>${{esc(r.district)}} · ${{esc(r.indicator)}}</h4><p>${{esc(r.sector)}} · ${{esc(r.period)}} · ${{esc(r.source)}}</p></span><span class="chip ${{statusClass(r.status)}}">${{statusText(r.status)}} · ${{fmt(r.execution_pct,'%')}}</span></button>`; }}
    function districtCard(d) {{ return `<button class="cardbtn ${{state.district===d.name?'active':''}}" data-district="${{esc(d.name)}}"><h4>${{esc(d.name)}}</h4><p>${{esc(d.owner)}} </p><strong>${{d.debt?.attention || 0}}/${{d.debt?.monitoring_total || 0}}</strong><p>эътибор / жами қатор</p></button>`; }}
    function renderDistricts() {{
      const d=DATA.districts.find(x=>x.name===state.district)||DATA.districts[0]; const rows=DATA.monitoring.filter(r=>r.district===d.name);
      $('#view-districts').innerHTML = `<div class="workspace"><div class="surface"><div class="surface-head"><div><h3>Туман/шаҳар</h3><p>${{DATA.districts.length}} ҳудуд.</p></div></div><div class="list">${{DATA.districts.map(districtCard).join('')}}</div></div><div class="profile"><div class="profile-head"><div><h3>${{esc(d.name)}}</h3><p>${{esc(d.owner)}} · full Andijan pilot data.</p></div><div class="debt"><strong>${{d.debt?.attention || 0}}/${{d.debt?.monitoring_total || 0}}</strong><span>эътибор / жами</span></div></div><div class="mini-kpis">${{districtMiniKpis(d).join('')}}</div><div class="surface"><div class="surface-head"><div><h3>Ҳудуд мониторинг қаторлари</h3><p>${{rows.length}} source-linked rows.</p></div></div><div class="table-wrap"><table><thead>${{monitorHead()}}</thead><tbody>${{rows.map(monitorRow).join('')}}</tbody></table></div></div></div></div>`;
      bindDistricts(); bindRows();
    }}
    function districtMiniKpis(d) {{
      const x=d.data||{{}}; return [
        ['Саноат',fmt(x.industry?.h1_value),'млрд сўм H1'],
        ['Инвестиция',fmt(x.foreign_investment?.q1_actual),'млн $ Q1'],
        ['Экспорт',fmt(x.export?.h1_expected),'минг $ H1'],
        ['Бюджет',fmt(x.budget?.h1_expected),'млрд H1'],
        ['Ишсизлик',fmt(x.employment?.unemployment_h1,'%'),'H1'],
        ['Камбағаллик',fmt(x.employment?.poverty_h1,'%'),'H1']
      ].map(k=>`<div class="mini"><b>${{esc(k[0])}}</b><strong>${{esc(k[1])}}</strong><span>${{esc(k[2])}}</span></div>`);
    }}
    function renderMonitoring() {{
      const rows=filteredRows();
      $('#view-monitoring').innerHTML=`<div class="surface"><div class="surface-head"><div><h3>Мониторинг реестри</h3><p>${{rows.length}} қатор фильтрга тушди.</p></div></div><div class="table-wrap"><table><thead>${{monitorHead()}}</thead><tbody>${{rows.map(monitorRow).join('')}}</tbody></table></div></div>`; bindRows();
    }}
    function monitorHead() {{ return '<tr><th>ID</th><th>Ҳудуд</th><th>Йўналиш</th><th>Кўрсаткич</th><th>Давр</th><th>Мақсад</th><th>Факт</th><th>Ижро</th><th>Ҳолат</th><th>Манба</th></tr>'; }}
    function monitorRow(r) {{ return `<tr class="action" data-row="${{r.id}}" tabindex="0" role="button"><td><b>${{esc(r.id)}}</b></td><td>${{esc(r.district)}}</td><td>${{esc(r.sector)}}</td><td>${{esc(r.indicator)}}</td><td>${{esc(r.period)}}</td><td>${{fmt(r.target)}} ${{esc(r.unit||'')}}</td><td>${{fmt(r.actual)}} ${{esc(r.unit||'')}}</td><td><b>${{fmt(r.execution_pct,'%')}}</b></td><td><span class="chip ${{statusClass(r.status)}}">${{statusText(r.status)}}</span></td><td>${{esc(r.source)}}</td></tr>`; }}
    function renderTasks() {{ $('#view-tasks').innerHTML=`<div class="surface"><div class="surface-head"><div><h3>Кафолат хати топшириқ номзодлари</h3><p>${{DATA.tasks.length}} матн қаторлари автоматик ажратилди; кейин масъул ва муддат билан тозаланади.</p></div></div><div class="pad">${{DATA.tasks.map((t,i)=>`<div class="rowbtn"><span class="num">${{i+1}}</span><span><h4>${{esc(t.title)}}</h4><p>${{esc(t.sector)}} · ${{esc(t.source)}}</p></span><span class="chip grey">номзод</span></div>`).join('')}}</div></div>`; }}
    function renderSources() {{ $('#view-sources').innerHTML=`<div class="surface"><div class="surface-head"><div><h3>Манбалар реестри</h3><p>${{DATA.sources.length}} файл/варақ қайди.</p></div></div><div class="table-wrap"><table><thead><tr><th>Файл</th><th>Тур</th><th>Варақ</th><th>Қатор</th><th>Устун</th><th>Тўлдирилган қатор</th></tr></thead><tbody>${{DATA.sources.map(s=>`<tr><td><b>${{esc(s.file)}}</b></td><td>${{esc(s.type)}}</td><td>${{esc(s.sheet)}}</td><td>${{esc(s.rows)}}</td><td>${{esc(s.cols)}}</td><td>${{esc(s.nonempty_rows)}}</td></tr>`).join('')}}</tbody></table></div></div>`; }}
    function renderQuality() {{ $('#view-quality').innerHTML=`<div class="surface"><div class="surface-head"><div><h3>Маълумот сифати текширувлари</h3><p>Импортда автоматик текшириладиган масалалар.</p></div></div><div class="pad">${{DATA.data_quality.map((q,i)=>`<div class="rowbtn"><span class="num">${{i+1}}</span><span><h4>${{esc(q.issue)}}</h4><p>${{esc(q.detail)}}</p></span><span class="chip ${{q.severity==='high'?'red':q.severity==='medium'?'amber':'grey'}}">${{esc(q.severity)}}</span></div>`).join('')}}</div></div>`; }}
    function bindDistricts() {{ $$('[data-district]').forEach(b=>b.addEventListener('click',()=>{{state.district=b.dataset.district; state.view='districts'; render();}})); }}
    function bindRows() {{ $$('[data-row]').forEach(el=>{{ const open=()=>{{ const r=DATA.monitoring.find(x=>x.id===el.dataset.row); if(!r) return; openDrawer(r); }}; el.addEventListener('click',open); el.addEventListener('keydown',e=>{{if(e.key==='Enter'||e.key===' '){{e.preventDefault();open();}}}}); }}); }}
    function openDrawer(r) {{ $('#drawerTitle').textContent=r.indicator; $('#drawerSub').textContent=`${{r.district}} · ${{r.period}}`; $('#drawerBody').innerHTML=`<div class="drawer-card"><dl class="kv"><dt>Ҳолат</dt><dd><span class="chip ${{statusClass(r.status)}}">${{statusText(r.status)}} · ${{fmt(r.execution_pct,'%')}}</span></dd><dt>Мақсад</dt><dd>${{fmt(r.target)}} ${{esc(r.unit||'')}}</dd><dt>Факт</dt><dd>${{fmt(r.actual)}} ${{esc(r.unit||'')}}</dd><dt>Йўналиш</dt><dd>${{esc(r.sector)}}</dd><dt>Манба</dt><dd>${{esc(r.source)}}</dd><dt>Далил</dt><dd>${{esc(r.evidence)}}</dd></dl></div><div class="drawer-card"><h3>Импорт изи</h3><p>Қатор Excel/DOCX манба номи ва sheet билан боғланган. Production босқичида row/column range ҳам сақланади.</p></div>`; $('#drawer').classList.add('open'); $('#mask').classList.add('open'); $('#drawer').setAttribute('aria-hidden','false'); }}
    function closeDrawer() {{ $('#drawer').classList.remove('open'); $('#mask').classList.remove('open'); $('#drawer').setAttribute('aria-hidden','true'); }}
    function exportRows() {{ const header=['ID','Ҳудуд','Йўналиш','Кўрсаткич','Давр','Мақсад','Факт','Ижро','Ҳолат','Манба']; const csv=[header,...filteredRows().map(r=>[r.id,r.district,r.sector,r.indicator,r.period,r.target,r.actual,r.execution_pct,statusText(r.status),r.source])].map(row=>row.map(c=>`"${{String(c??'').replace(/"/g,'""')}}"`).join(';')).join('\\r\\n'); const blob=new Blob(['\\ufeff'+csv],{{type:'text/csv;charset=utf-8'}}); const a=document.createElement('a'); a.href=URL.createObjectURL(blob); a.download='andijon_full_pilot_monitoring.csv'; a.click(); URL.revokeObjectURL(a.href); }}
    init();
  </script>
</body>
</html>
"""


def write_html(data):
    data_json = json.dumps(data, ensure_ascii=False)
    OUT_HTML.write_text(html_template(data_json), encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = build_dataset()
    OUT_JSON.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    write_audit(data)
    write_html(data)
    print(json.dumps({
        "json": str(OUT_JSON.relative_to(ROOT)),
        "audit": str(OUT_AUDIT.relative_to(ROOT)),
        "html": str(OUT_HTML.relative_to(ROOT)),
        "monitoring_rows": len(data["monitoring"]),
        "tasks": len(data["tasks"]),
        "districts": len(data["districts"]),
        "sources": len(data["sources"]),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
